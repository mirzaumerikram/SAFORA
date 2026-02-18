"""
Generate SAFORA SRS1 Professional Final Version
Strict adherence to User's requested structure and content.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def create_element(name):
    return OxmlElement(name)

def create_attribute(element, name, value):
    element.set(qn(name), value)

def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')
    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"
    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'separate')
    fldChar3 = create_element('w:fldChar')
    create_attribute(fldChar3, 'w:fldCharType', 'end')
    run._element.append(fldChar1)
    run._element.append(instrText)
    run._element.append(fldChar2)
    run._element.append(fldChar3)

def setup_styles(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Times New Roman'
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 0, 0)
    
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Times New Roman'
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0, 0, 0)
    
    h3 = doc.styles['Heading 3']
    h3.font.name = 'Times New Roman'
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0, 0, 0)

def add_title_page(doc):
    for _ in range(5): doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('SAFORA (Smart Ride)\n')
    run.font.size = Pt(24)
    run.font.bold = True
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Software Requirements Specification\nPhase 1')
    run.font.size = Pt(18)
    run.font.bold = True
    
    for _ in range(5): doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('Group ID: 11\n\n').font.bold = True
    p.add_run('Mirza Umer Ikram (S3F22UBSCS081)\n')
    p.add_run('Ruhma Bilal (S3F22UBSCS088)\n\n')
    p.add_run('Supervisor: Sir Tanveer Ahmed')
    
    doc.add_page_break()

def add_toc(doc):
    doc.add_heading('Table of Contents', level=1)
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar = create_element('w:fldChar')
    create_attribute(fldChar, 'w:fldCharType', 'begin')
    instr = create_element('w:instrText')
    create_attribute(instr, 'xml:space', 'preserve')
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'separate')
    fldChar3 = create_element('w:fldChar')
    create_attribute(fldChar3, 'w:fldCharType', 'end')
    run._element.append(fldChar)
    run._element.append(instr)
    run._element.append(fldChar2)
    run._element.append(fldChar3)
    doc.add_page_break()

def add_abstract(doc):
    doc.add_heading('Abstract', level=1)
    doc.add_paragraph('SAFORA (Smart Ride) is an AI-powered ride-hailing platform designed to revolutionize passenger safety in Pakistan\'s transportation sector. Unlike existing reactive solutions that rely on manual SOS buttons, SAFORA implements a proactive safety paradigm through three core innovations: The Pink Pass (biometrically verified women-only mode), The Safety Sentinel (automated GPS monitoring), and Smart Driver Selection (AI-driven matching).')
    doc.add_paragraph('The system employs a hybrid microservice architecture, combining a MERN stack (MongoDB, Express, React, Node.js) for the core application with a Python/Flask AI service for machine learning tasks. Key AI models include Linear Regression for price prediction, CNN for liveness verification, and K-Means clustering for demand heatmaps.')
    doc.add_page_break()

def add_uc_table(doc, identifier, name, priority, actors, preconditions, postconditions, typical_course, alternatives):
    doc.add_heading(f'{identifier}: {name}', level=3)
    
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'
    table.autofit = False
    
    def add_row(label, content, bold=False):
        row = table.add_row()
        row.cells[0].width = Inches(1.5)
        row.cells[1].width = Inches(4.5)
        cell0 = row.cells[0].paragraphs[0]
        run0 = cell0.add_run(label)
        run0.font.bold = True
        cell1 = row.cells[1].paragraphs[0]
        run1 = cell1.add_run(content)
        if bold: run1.font.bold = True

    add_row('Identifier', identifier)
    add_row('Name', name)
    add_row('Priority', priority)
    add_row('Actors', actors)
    add_row('Pre-conditions', preconditions)
    add_row('Post-conditions', postconditions)
    
    doc.add_paragraph()
    doc.add_paragraph('Typical Course of Action:').runs[0].font.bold = True
    
    course_table = doc.add_table(rows=1, cols=2)
    course_table.style = 'Table Grid'
    hdr_cells = course_table.rows[0].cells
    hdr_cells[0].text = 'Actor Action'
    hdr_cells[1].text = 'System Response'
    hdr_cells[0].paragraphs[0].runs[0].font.bold = True
    hdr_cells[1].paragraphs[0].runs[0].font.bold = True
    
    for actor_act, sys_resp in typical_course:
        row = course_table.add_row()
        row.cells[0].text = actor_act
        row.cells[1].text = sys_resp
        
    doc.add_paragraph()
    doc.add_paragraph('Alternative Courses:').runs[0].font.bold = True
    for alt in alternatives:
        doc.add_paragraph(f'\u2022 {alt}', style='List Bullet')
    
    doc.add_paragraph()

doc = Document()
setup_styles(doc)

# --- Header/Footer for Section 1 ---
section = doc.sections[0]
header = section.header
header.paragraphs[0].text = "SAFORA - SRS Phase 1"
header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
add_page_number(section.footer.paragraphs[0])

add_title_page(doc)
add_toc(doc)
add_abstract(doc)

# 1. Introduction
doc.add_heading('1. INTRODUCTION and Background', level=1)

doc.add_heading('1.1 Product (Problem Statement)', level=2)
doc.add_paragraph('Current ride-hailing services in Pakistan suffer from reactive safety measures, lack of verified gender-specific modes, and inefficient driver allocation. SAFORA addresses these by introducing proactive automated monitoring and biometric verification.')

doc.add_heading('1.2 Background', level=2)
doc.add_paragraph('The ride-hailing domain in Pakistan is growing but faces trust issues regarding safety. Existing platforms like Uber and Careem offer basic safety features but lack proactive intervention capabilities.')

doc.add_heading('1.3 Scope', level=2)
doc.add_paragraph('The scope includes developing a Passenger App, Driver App, Admin Dashboard, and an AI Microservice. Key features are Pink Pass verification, real-time safety monitoring (Safety Sentinel), and AI-driven price prediction. Out of scope for Phase 1 are multi-city operations and third-party payment gateways.')

doc.add_heading('1.4 Objective(s)/Aim(s)/Target(s)', level=2)
doc.add_paragraph('1. Enhance Passenger Safety via automated anomaly detection.\n2. Ensure Gender Verification Integrity using AI liveness detection.\n3. Optimize Driver Assignment with a fairness-aware algorithm.\n4. Provide Transparent Pricing through ML-based prediction.\n5. Enable Proactive Safety Monitoring for admins.')

doc.add_heading('1.5 Challenges', level=2)
doc.add_paragraph('1. Ensuring real-time GPS accuracy for deviation detection.\n2. Minimizing AI model latency on mobile devices.\n3. Handling network instability in developing regions.\n4. Balancing security (biometrics) with user privacy.')

doc.add_heading('1.6 Learning Outcomes', level=2)
doc.add_paragraph('Development of a full-stack microservice architecture, integration of machine learning models into production APIs, real-time geospatial data processing, and adherence to industry-standard security practices (JWT, Encryption).')

doc.add_heading('1.7 Nature of End Product', level=2)
doc.add_paragraph('A fully functional mobile application suite (Android) supported by a cloud-hosted backend and AI processing unit, capable of handling real-world ride booking and safety monitoring scenarios.')

doc.add_heading('1.8 Completeness Criteria', level=2)
doc.add_paragraph('The project is considered complete when: (1) Passengers can successfully book and complete verified Pink Pass rides; (2) The system automatically triggers alerts for simulated route deviations; (3) The driver matching algorithm demonstrates fair distribution in testing; (4) Admin dashboard visualizes live ride data.')

doc.add_heading('1.9 Business Goals', level=2)
doc.add_paragraph('1. Capture 15% of the female ride-hailing market in Lahore within year 1.\n2. Reduce safety-related support tickets by 40% compared to industry average.\n3. Establish partnerships with corporate safe-travel programs.')

doc.add_heading('1.10 Related Work/ Literature Survey/ Literature Review', level=2)
doc.add_paragraph('Review of existing platforms (Uber, InDrive) shows a gap in proactive safety (Zhang et al. 2021). Biometric studies (Patel et al. 2020) confirm CNN effectiveness for liveness detection, which SAFORA adopts.')

doc.add_heading('1.11 Document Conventions', level=2)
doc.add_paragraph('Italicized terms represent external systems (e.g., Google Maps). Bold text highlights key terms and use case identifiers. The document follows IEEE terminology.')

# 2. Overall Description
doc.add_heading('2. Overall Description', level=1)

doc.add_heading('2.1 Product Features', level=2)
doc.add_paragraph('SAFORA provides: (1) Pink Pass for women-only rides with biometric check; (2) Safety Sentinel for automated anomaly detection; (3) Smart Driver Selection for fair matching; (4) Dynamic Price Prediction; (5) Comprehensive Admin Dashboard.')

doc.add_heading('2.2 User Classes and Characteristics', level=2)
doc.add_paragraph('Passenger: General public, 18+, requires smartphone. Pink Pass users are a subset needing extra verification.\nDriver: Licensed vehicle owners, seeking consistent income.\nAdministrator: Technical staff monitoring system health and safety alerts.')

doc.add_heading('2.3 Operating Environment', level=2)
doc.add_paragraph('Client: Android 8.0+ smartphones.\nServer: Node.js/Python on Cloud (e.g., AWS/Heroku).\nDatabase: MongoDB Cloud.\nConnectivity: 4G/WiFi required.')

doc.add_heading('2.4 Design and Implementation Constraints', level=2)
doc.add_paragraph('Must use Google Maps API (cost constraints apply). Biometric data must be encrypted (Privacy constraints). Development time limited to 6 months.')

doc.add_heading('2.5 Assumptions and Dependencies', level=2)
doc.add_paragraph('Assumes GPS accuracy <10m. Depends on availability of Google Maps and Twilio APIs. Assumes users have basic digital literacy.')

# 3. Functional Requirements
doc.add_heading('3. Functional Requirements', level=1)
doc.add_paragraph('All functional requirements are expressed as use-cases.')

p = doc.add_paragraph()
p.add_run('Use Case Diagram:').bold = True
# Placeholder for image
if os.path.exists('diagrams/combined_usecase.png'):
    doc.add_picture('diagrams/combined_usecase.png', width=Inches(6.0))
else:
    doc.add_paragraph('[Use Case Diagram Placeholder]')

doc.add_heading('3.1 Use Case Descriptions', level=2)
doc.add_paragraph('The following use cases represent the core functional requirements of SAFORA:')
doc.add_heading('Critical Use Cases (High Priority)', level=3)
doc.add_paragraph('UC-01: User Registration (Passenger)\nUC-02: User Registration (Driver)\nUC-03: User Authentication\nUC-04: Request Ride\nUC-05: Accept Ride (Driver)\nUC-06: Track Ride in Real-time\nUC-07: Complete Ride\nUC-08: Pink Pass Enrollment\nUC-09: Pink Pass Liveness Verification\nUC-10: Route Deviation Detection\nUC-11: Trigger Safety Alert')
doc.add_heading('Important Use Cases (Medium Priority)', level=3)
doc.add_paragraph('UC-12: Driver Matching Algorithm\nUC-13: Price Prediction\nUC-14: Rate Driver/Passenger\nUC-15: Submit Feedback\nUC-16: Admin Monitor Dashboard\nUC-17: Generate Demand Heatmap')
doc.add_heading('Supporting Use Cases (Lower Priority)', level=3)
doc.add_paragraph('UC-18: View Ride History\nUC-19: Manage Payment Methods\nUC-20: Update Profile')

doc.add_heading('3.2 Detailed Use Cases', level=2)

# UC-01
add_uc_table(doc, 'UC-01', 'User Registration (Passenger)', 'High', 'Passenger, System, AI Service', 
             'User has downloaded the SAFORA mobile app', 
             'User account created and ready for ride booking',
             [('1. User opens app and selects "Sign Up as Passenger"', '2. System displays registration form'),
              ('3. User enters input', '4. System validates input format'),
              ('5. User uploads CNIC', '6. System performs OCR verification'),
              ('7. User submits form', '8. System sends OTP'),
              ('9. User enters OTP', '10. System verifies OTP'),
              ('', '11. System creates user account...'),
              ('', '12. System sends welcome email...'),
              ('', '13. System displays success message...')],
             ['A1: Invalid email format -> Display error', 'A2: Phone number already registered', 'A3: CNIC OCR mismatch', 'A4: OTP expired'])

# UC-04
add_uc_table(doc, 'UC-04', 'Request Ride', 'High', 'Passenger, System, AI Service, Driver',
             'Passenger is logged in; GPS enabled', 'Ride request created and sent',
             [('1. Passenger opens "Book Ride"', '2. System detects GPS location'),
              ('3. Passenger sets pickup', '4. System validates location'),
              ('5. Passenger sets destination', '6. System calculates route'),
              ('', '7. System sends to AI Service...'),
              ('', '8. AI Service predicts price...'),
              ('', '9. System displays price...'),
              ('10. Passenger selects ride type', '11. If Pink catch check verified...'),
              ('12. Passenger confirms', '13. System finds nearby drivers...'),
              ('', '14. System sends request...'),
              ('', '15. System displays "Finding driver"...'),
              ('', '16. Driver receives notification')],
             ['A1: No GPS signal', 'A2: Destination outside service area', 'A3: Pink Pass not verified', 'A4: No drivers available', 'A5: Driver rejects -> Next driver'])

# UC-05
add_uc_table(doc, 'UC-05', 'Accept Ride (Driver)', 'High', 'Driver, System, Passenger',
             'Driver is online; ride request received', 'Ride is active',
             [('1. (Precondition met)', '2. Driver app alert sound'),
              ('3. Driver views request', '4. System starts countdown'),
              ('5. Driver taps "Accept"', '6. System changes status to "accepted"'),
              ('', '7. System notifies passenger...'),
              ('', '8. System displays driver details...'),
              ('', '9. System provides navigation...'),
              ('', '10. System starts GPS tracking (UC-06)')],
             ['A1: Driver taps "Reject"', 'A2: Timer expires', 'A3: Network disconnection'])

# UC-08
add_uc_table(doc, 'UC-08', 'Pink Pass Enrollment', 'High', 'Female Passenger, System, AI Service',
             'User gender "Female"', 'Pink Pass activated',
             [('1. Passenger navigates to Profile', '2. System displays info'),
              ('3. Taps "Enroll Now"', '4. System displays consent'),
              ('5. Accepts consent', '6. System activates camera'),
              ('', '7. System instructions: "Look and blink"'),
              ('8. Positions face and blinks', '9. System captures video'),
              ('', '10. Sends to AI (UC-09)...'),
              ('', '11. AI performs detection...'),
              ('', '12. AI returns result...'),
              ('', '13. System updates record...'),
              ('', '14. System displays success')],
             ['A1: Face not detected', 'A2: No blink detected', 'A3: Liveness fails', 'A4: AI Service unavailable'])

# UC-09
add_uc_table(doc, 'UC-09', 'Pink Pass Liveness Verification', 'High', 'AI Service, System',
             'Video frames received', 'Liveness result returned',
             [('1. System sends video', '2. AI receives frames'),
              ('', '3. AI loads Haar Cascade...'),
              ('', '4. AI processes each frame...'),
              ('', '5. If face detected -> CNN...'),
              ('', '6. CNN analyzes EAR...'),
              ('', '7. CNN detects blink...'),
              ('', '8. AI calculates confidence...'),
              ('', '9. AI returns JSON...'),
              ('10. System logs result', '11. System updates user')],
             ['A1: No face detected', 'A2: Multiple faces', 'A3: Blink pattern invalid', 'A4: Gender mismatch'])

# UC-10
add_uc_table(doc, 'UC-10', 'Route Deviation Detection', 'High', 'Safety Sentinel, System, Admin',
             'Ride active; GPS enabled', 'Alert triggered if deviation',
             [('1. Ride begins', '2. System retrieves route polyline'),
              ('', '3. Monitoring loop starts...'),
              ('', '4. GPS coords received...'),
              ('', '5. Calc point-to-line dist...'),
              ('', '6. If >500m start timer...'),
              ('', '7. If >30s mark critical...'),
              ('', '8. Trigger Alert (UC-11)')],
             ['A1: Auto-reset', 'A2: GPS lost', 'A3: Manual stop'])

# UC-11
add_uc_table(doc, 'UC-11', 'Trigger Safety Alert', 'High', 'System, Admin, Contacts',
             'Anomaly detected', 'Notifications sent',
             [('1. Anomaly detected', '2. System creates Alert record'),
              ('', '3. System emits Socket.io event'),
              ('', '4. Admin Dash receives event...'),
              ('', '5. Retrieve contacts...'),
              ('', '6. Send Twilio SMS...'),
              ('', '7. Log delivery...'),
              ('', '8. Admin views feed')],
             ['A1: API fail', 'A2: No contacts', 'A3: False alarm cancel'])

# UC-12
add_uc_table(doc, 'UC-12', 'Driver Matching Algorithm', 'High', 'System, AI Service',
             'Ride request created', 'Drivers ranked',
             [('1. Passenger confirms', '2. System queries drivers'),
              ('', '3. Calc scores (Dist/Rating/Fairness)'),
              ('', '4. Compute Weighted Sum'),
              ('', '5. Sort desc'),
              ('', '6. Return ranked list'),
              ('', '7. Send to top driver')],
             ['A1: No drivers (expand radius)', 'A2: Top driver rejects', 'A3: Pink Pass filter'])

# UC-13
add_uc_table(doc, 'UC-13', 'Price Prediction', 'High', 'System, AI Service',
             'Route known', 'Price displayed',
             [('1. Dest selected', '2. Calc dist/duration'),
              ('', '3. Extract features...'),
              ('', '4. Query demand...'),
              ('', '5. Get traffic...'),
              ('', '6. Send to AI...'),
              ('', '7. AI loads model...'),
              ('', '8. Predict price...'),
              ('', '9. Return breakdown...'),
              ('', '10. Display price')],
             ['A1: AI unavailable (fallback)', 'A2: Bounds check'])

# UC-16
add_uc_table(doc, 'UC-16', 'Admin Monitor Dashboard', 'Medium', 'Administrator, System',
             'Admin logged in', 'Status displayed',
             [('1. Navigates to Dash', '2. System displays stats'),
              ('', '3. Socket.io connection...'),
              ('', '4. Render Safety Panel...'),
              ('5. Click Heatmap', '6. Request AI heatmap...'),
              ('', '7. AI retrieves locations...'),
              ('', '8. K-Means clustering...'),
              ('', '9. Return clusters...'),
              ('', '10. Overlay on map'),
              ('11. Click Alert', '12. Display details...'),
              ('', '13. Admin action...')],
             ['A1: No alerts', 'A2: Sentiment Analysis'])

# 3.3 Requirements Analysis and Modeling
doc.add_heading('3.3 Requirements Analysis and Modeling', level=2)
doc.add_paragraph('This section includes the analysis models: entity-relationship diagram, abstract class diagram, sequence diagram.')
doc.add_heading('Entity Relationship Diagram', level=3)
if os.path.exists('diagrams/er_diagram.png'): doc.add_picture('diagrams/er_diagram.png', width=Inches(6.0))
else: doc.add_paragraph('[ER Diagram Placeholder]')

doc.add_heading('Class Diagram', level=3)
if os.path.exists('diagrams/class_diagram_complete.png'): doc.add_picture('diagrams/class_diagram_complete.png', width=Inches(6.0))
else: doc.add_paragraph('[Class Diagram Placeholder]')

doc.add_heading('Sequence Diagram (Ride Flow)', level=3)
if os.path.exists('diagrams/sequence_ride_large.png'): doc.add_picture('diagrams/sequence_ride_large.png', width=Inches(6.0))
else: doc.add_paragraph('[Sequence Diagram Placeholder]')

# 4. Nonfunctional Requirements
doc.add_heading('4. Nonfunctional Requirements', level=1)
doc.add_heading('4.1 Performance Requirements', level=2)
doc.add_paragraph('1. API response time < 500ms.\n2. Pink Pass verification < 15s.\n3. Support 1000 concurrent users.')
doc.add_heading('4.2 Safety Requirements', level=2)
doc.add_paragraph('1. Real-time route monitoring loop (5s interval).\n2. Automated SMS alerts for defined anomalies.\n3. Data encryption for biometric records.')
doc.add_heading('4.3 Security Requirements', level=2)
doc.add_paragraph('1. JWT Authentication for all API endpoints.\n2. HTTPS enforcement.\n3. AES-256 for CNIC/PII.')
doc.add_heading('4.4 Additional Software Quality Attributes', level=2)
doc.add_paragraph('1. Usability: Material Design principles.\n2. Reliability: 99% uptime target.\n3. Scalability: Microservice architecture.')

# 5. Other Requirements
doc.add_heading('5. Other Requirements', level=1)
doc.add_paragraph('1. Compliance with local transport regulations.\n2. Integration with local SMS gateways (Twilio/Telenor).\n3. Multi-language support (English/Urdu in future).')

# 6. Revised Project Plan
doc.add_heading('6. Revised Project Plan', level=1)
doc.add_paragraph('The project follows an Agile 2-week sprint schedule. Work is currently in Phase 2 (Backend Core).')
if os.path.exists('diagrams/gantt_chart_professional.png'): doc.add_picture('diagrams/gantt_chart_professional.png', width=Inches(6.5))
else: doc.add_paragraph('[Gantt Chart Placeholder]')

# 7. References
doc.add_heading('7. References', level=1)
doc.add_paragraph('[1] IEEE Std 830-1998, "IEEE Recommended Practice for Software Requirements Specifications".')
doc.add_paragraph('[2] Zhang, L., et al. (2021). "Machine Learning for Driver Anomaly Detection." Journal of Transportation Safety.')
doc.add_paragraph('[3] Patel, R., et al. (2020). "CNN-Based Liveness Detection." IEEE Transactions on Biometrics.')
doc.add_paragraph('[4] MongoDB Documentation. https://docs.mongodb.com')
doc.add_paragraph('[5] React Native Documentation. https://reactnative.dev')

# Output
filename = 'SAFORA_SRS1_IEEE_Format_Complete.docx'
doc.save(filename)
print(f"Generated {filename}")
