"""
Convert SAFORA Phase 1 Markdown to MS Word Document with proper formatting
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def add_page_break(doc):
    """Add a page break"""
    doc.add_page_break()

def set_cell_border(cell, **kwargs):
    """Set cell borders"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = OxmlElement(tag)
            element.set(qn('w:val'), 'single')
            element.set(qn('w:sz'), '4')
            element.set(qn('w:space'), '0')
            element.set(qn('w:color'), '000000')
            tcPr.append(element)

def parse_markdown_to_word(md_file, output_file):
    """Parse markdown and create Word document"""
    
    # Create document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Skip empty lines at start
        if not line:
            i += 1
            continue
        
        # Title Page (First heading)
        if i < 20 and line.startswith('# '):
            title = line[2:].strip()
            p = doc.add_paragraph(title)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].font.size = Pt(18)
            p.runs[0].font.bold = True
            i += 1
            continue
        
        # H2 - Subtitle
        if line.startswith('## ') and i < 30:
            subtitle = line[3:].strip()
            p = doc.add_paragraph(subtitle)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].font.size = Pt(14)
            p.runs[0].font.bold = True
            i += 1
            continue
        
        # H3 - Small subtitle
        if line.startswith('### ') and i < 30:
            subtitle = line[4:].strip()
            p = doc.add_paragraph(subtitle)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].font.size = Pt(12)
            p.runs[0].font.bold = True
            i += 1
            continue
        
        # Horizontal rule or page markers - add page break
        if line.strip() == '---':
            add_page_break(doc)
            i += 1
            continue
        
        # Handle tables
        if line.startswith('|') and '|' in line:
            # Collect table rows
            table_rows = []
            while i < len(lines) and lines[i].startswith('|'):
                table_rows.append(lines[i])
                i += 1
            
            # Parse table
            if len(table_rows) >= 2:
                # First row is header
                header = [cell.strip() for cell in table_rows[0].split('|')[1:-1]]
                
                # Second row is separator (skip)
                # Remaining rows are data
                data_rows = []
                for row in table_rows[2:]:
                    cells = [cell.strip() for cell in row.split('|')[1:-1]]
                    if any(cells):  # Skip empty rows
                        data_rows.append(cells)
                
                if header and data_rows:
                    # Create table
                    table = doc.add_table(rows=1 + len(data_rows), cols=len(header))
                    table.style = 'Light Grid Accent 1'
                    
                    # Header row
                    for idx, cell_text in enumerate(header):
                        cell = table.rows[0].cells[idx]
                        cell.text = cell_text
                        cell.paragraphs[0].runs[0].font.bold = True
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Data rows
                    for row_idx, row_data in enumerate(data_rows):
                        for col_idx, cell_text in enumerate(row_data):
                            if col_idx < len(table.rows[row_idx + 1].cells):
                                table.rows[row_idx + 1].cells[col_idx].text = cell_text
                    
                    doc.add_paragraph()  # Add space after table
            continue
        
        # Headings
        if line.startswith('# '):
            heading = line[2:].strip()
            p = doc.add_heading(heading, level=1)
            p.runs[0].font.size = Pt(16)
            p.runs[0].font.bold = True
            i += 1
            continue
        
        if line.startswith('## '):
            heading = line[3:].strip()
            p = doc.add_heading(heading, level=2)
            p.runs[0].font.size = Pt(14)
            p.runs[0].font.bold = True
            i += 1
            continue
        
        if line.startswith('### '):
            heading = line[4:].strip()
            p = doc.add_heading(heading, level=3)
            p.runs[0].font.size = Pt(13)
            p.runs[0].font.bold = True
            i += 1
            continue
        
        if line.startswith('#### '):
            heading = line[5:].strip()
            p = doc.add_heading(heading, level=4)
            p.runs[0].font.size = Pt(12)
            p.runs[0].font.bold = True
            i += 1
            continue
        
        # Blockquotes (Notes)
        if line.startswith('> '):
            note_lines = []
            while i < len(lines) and lines[i].startswith('> '):
                note_lines.append(lines[i][2:].strip())
                i += 1
            
            p = doc.add_paragraph(' '.join(note_lines))
            p.style = 'Intense Quote'
            continue
        
        # Bullet lists
        if line.startswith('- ') or line.startswith('* '):
            bullet_text = line[2:].strip()
            # Check for checkboxes
            if bullet_text.startswith('[ ] '):
                bullet_text = '☐ ' + bullet_text[4:]
            elif bullet_text.startswith('[x] '):
                bullet_text = '☑ ' + bullet_text[4:]
            elif bullet_text.startswith('[/] '):
                bullet_text = '◐ ' + bullet_text[4:]
            
            p = doc.add_paragraph(bullet_text, style='List Bullet')
            i += 1
            continue
        
        # Numbered lists
        if re.match(r'^\d+\. ', line):
            list_text = re.sub(r'^\d+\. ', '', line)
            p = doc.add_paragraph(list_text, style='List Number')
            i += 1
            continue
        
        # Code blocks
        if line.startswith('```'):
            code_lines = []
            i += 1  # Skip opening ```
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1  # Skip closing ```
            
            if code_lines:
                p = doc.add_paragraph('\n'.join(code_lines))
                p.style = 'No Spacing'
                p.runs[0].font.name = 'Courier New'
                p.runs[0].font.size = Pt(10)
            continue
        
        # Regular paragraph
        if line.strip():
            # Handle bold **text**
            line = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
            # Handle italic *text*
            line = re.sub(r'\*(.*?)\*', r'\1', line)
            # Handle inline code `text`
            line = re.sub(r'`(.*?)`', r'\1', line)
            
            p = doc.add_paragraph(line)
        
        i += 1
    
    # Save document
    doc.save(output_file)
    print(f"Successfully converted to {output_file}")

if __name__ == '__main__':
    md_file = r'C:\Users\mzees\OneDrive\Desktop\SAFORA\SAFORA_Phase1_SRS_Condensed.md'
    output_file = r'C:\Users\mzees\OneDrive\Desktop\SAFORA\SAFORA_Phase1_SRS_Condensed.docx'
    
    parse_markdown_to_word(md_file, output_file)
