"""
SAFORA SRS2 Document Generator with Mermaid Diagrams
Generates comprehensive SRS2 with all UML diagrams
"""

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create document
doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# ==================== TITLE PAGE ====================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run('SAFORA (Smart Ride)\n')
title_run.font.size = Pt(20)
title_run.font.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.add_run('Software Requirements Specification\nPhase 2 - System Design & Diagrams')
subtitle_run.font.size = Pt(18)
subtitle_run.font.bold = True

doc.add_paragraph()
doc.add_paragraph()

# University info
uni = doc.add_paragraph()
uni.alignment = WD_ALIGN_PARAGRAPH.CENTER
uni_run = uni.add_run('University of Central Punjab\n')
uni_run.font.bold = True
uni_run.font.size = Pt(14)
uni.add_run('Faculty of Information Technology\n')
uni.add_run('Department of Computer Science')

doc.add_paragraph()
doc.add_paragraph()

# Project details
details = doc.add_paragraph()
details.alignment = WD_ALIGN_PARAGRAPH.CENTER
details.add_run('Project Title: ').font.bold = True
details.add_run('SAFORA - Smart Ride\n\n')
details.add_run('Project Advisor: ').font.bold = True
details.add_run('Tanveer Ahmed\n\n')
details.add_run('Group ID: ').font.bold = True
details.add_run('11\n\n')

team = details.add_run('Team Members:\n')
team.font.bold = True
details.add_run('Mirza Umer Ikram (S3F22UBSCS081)\n')
details.add_run('Ruhma Bilal (S3F22UBSCS088)\n\n')

date_run = details.add_run('Date: January 20, 2026')
date_run.font.bold = True

doc.add_page_break()

# ==================== TABLE OF CONTENTS ====================
doc.add_heading('Table of Contents', level=1)

toc_items = [
    '1. Use Case Diagrams',
    '   1.1 Passenger Use Cases',
    '   1.2 Driver Use Cases',
    '   1.3 Admin Use Cases',
    '2. Sequence Diagrams',
    '   2.1 User Registration Flow',
    '   2.2 Ride Request Flow',
    '   2.3 Pink Pass Verification Flow',
    '   2.4 Safety Alert Flow',
    '3. Class Diagram',
    '4. Entity Relationship Diagram',
    '5. Activity Diagrams',
    '   5.1 Ride Booking Process',
    '   5.2 Emergency Alert Process',
    '6. State Diagrams',
    '   6.1 Ride State Diagram',
    '   6.2 Driver State Diagram',
    '7. Component Diagram',
    '8. Deployment Diagram'
]

for item in toc_items:
    doc.add_paragraph(item)

doc.add_page_break()

# ==================== 1. USE CASE DIAGRAMS ====================
doc.add_heading('1. USE CASE DIAGRAMS', level=1)

doc.add_heading('1.1 Passenger Use Cases', level=2)
doc.add_paragraph('The following diagram shows all use cases available to passengers in the SAFORA system:')
doc.add_paragraph()

# Mermaid diagram description
p = doc.add_paragraph()
p.add_run('Diagram Description:\n').font.bold = True
doc.add_paragraph('Register Account - New users create passenger accounts', style='List Bullet')
doc.add_paragraph('Login - Authenticate using credentials', style='List Bullet')
doc.add_paragraph('Enroll in Pink Pass - Female passengers complete biometric verification', style='List Bullet')
doc.add_paragraph('Request Ride - Book a ride with pickup/dropoff locations', style='List Bullet')
doc.add_paragraph('Track Ride - Monitor real-time ride progress', style='List Bullet')
doc.add_paragraph('Rate Driver - Provide feedback after ride completion', style='List Bullet')
doc.add_paragraph('Trigger SOS - Activate emergency alert system', style='List Bullet')
doc.add_paragraph('View Ride History - Access past ride records', style='List Bullet')

doc.add_page_break()

doc.add_heading('1.2 Driver Use Cases', level=2)
doc.add_paragraph('The following diagram shows all use cases available to drivers in the SAFORA system:')
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Diagram Description:\n').font.bold = True
doc.add_paragraph('Register as Driver - Create driver account with license/vehicle info', style='List Bullet')
doc.add_paragraph('Login - Authenticate using credentials', style='List Bullet')
doc.add_paragraph('Update Availability - Toggle online/offline status', style='List Bullet')
doc.add_paragraph('Accept/Reject Ride - Respond to ride requests', style='List Bullet')
doc.add_paragraph('Navigate to Pickup - Use GPS to reach passenger location', style='List Bullet')
doc.add_paragraph('Start Ride - Begin trip after passenger pickup', style='List Bullet')
doc.add_paragraph('Complete Ride - End trip at destination', style='List Bullet')
doc.add_paragraph('View Earnings - Check income and ride statistics', style='List Bullet')

doc.add_page_break()

doc.add_heading('1.3 Admin Use Cases', level=2)
doc.add_paragraph('The following diagram shows all use cases available to administrators:')
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Diagram Description:\n').font.bold = True
doc.add_paragraph('Monitor Active Rides - Real-time view of all ongoing rides', style='List Bullet')
doc.add_paragraph('View Safety Alerts - Access automated safety notifications', style='List Bullet')
doc.add_paragraph('Manage Users - Add/remove/suspend user accounts', style='List Bullet')
doc.add_paragraph('Generate Reports - Create analytics and performance reports', style='List Bullet')
doc.add_paragraph('Analyze Heatmaps - View demand patterns using K-Means clustering', style='List Bullet')
doc.add_paragraph('Resolve Alerts - Respond to and close safety incidents', style='List Bullet')

doc.add_page_break()

# ==================== 2. SEQUENCE DIAGRAMS ====================
doc.add_heading('2. SEQUENCE DIAGRAMS', level=1)

doc.add_heading('2.1 User Registration Flow', level=2)
doc.add_paragraph('This sequence diagram illustrates the user registration process:')
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Flow Steps:\n').font.bold = True
doc.add_paragraph('1. User submits registration form with personal details', style='List Number')
doc.add_paragraph('2. Backend validates CNIC format and uniqueness', style='List Number')
doc.add_paragraph('3. Backend hashes password using bcrypt', style='List Number')
doc.add_paragraph('4. User record is created in MongoDB', style='List Number')
doc.add_paragraph('5. JWT token is generated for authentication', style='List Number')
doc.add_paragraph('6. Success response with token is returned to user', style='List Number')

doc.add_page_break()

doc.add_heading('2.2 Ride Request Flow', level=2)
doc.add_paragraph('This sequence diagram shows the complete ride request and matching process:')
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Flow Steps:\n').font.bold = True
doc.add_paragraph('1. Passenger selects pickup and dropoff locations', style='List Number')
doc.add_paragraph('2. Backend calculates route using Google Maps API', style='List Number')
doc.add_paragraph('3. Backend calls AI Service for price prediction', style='List Number')
doc.add_paragraph('4. AI Service returns estimated price based on distance, time, demand', style='List Number')
doc.add_paragraph('5. Backend queries available drivers within radius', style='List Number')
doc.add_paragraph('6. AI Service ranks drivers using weighted algorithm', style='List Number')
doc.add_paragraph('7. Ride request is sent to top-matched driver via Socket.io', style='List Number')
doc.add_paragraph('8. Driver accepts ride request', style='List Number')
doc.add_paragraph('9. Passenger receives driver details and ETA', style='List Number')

doc.add_page_break()

doc.add_heading('2.3 Pink Pass Verification Flow', level=2)
doc.add_paragraph('This sequence diagram illustrates the Pink Pass biometric verification process:')
doc.add_paragraph()

p= doc.add_paragraph()
p.add_run('Flow Steps:\n').font.bold = True
doc.add_paragraph('1. Female passenger uploads verification video (5-10 seconds)', style='List Number')
doc.add_paragraph('2. Backend uploads video to AI Service', style='List Number')
doc.add_paragraph('3. AI Service extracts frames from video', style='List Number')
doc.add_paragraph('4. Haar Cascade detects face in frames', style='List Number')
doc.add_paragraph('5. CNN model verifies liveness (blink detection)', style='List Number')
doc.add_paragraph('6. AI Service calculates confidence score', style='List Number')
doc.add_paragraph('7. Verification result (approved/rejected) is returned', style='List Number')
doc.add_paragraph('8. Backend updates user pinkPassVerified status', style='List Number')
doc.add_paragraph('9. Passenger receives confirmation notification', style='List Number')

doc.add_page_break()

doc.add_heading('2.4 Safety Alert Flow', level=2)
doc.add_paragraph('This sequence diagram shows the automated safety alert process:')
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Flow Steps:\n').font.bold = True
doc.add_paragraph('1. Driver location updates every 5 seconds during ride', style='List Number')
doc.add_paragraph('2. Safety Sentinel calculates point-to-line distance from planned route', style='List Number')
doc.add_paragraph('3. System detects deviation > 500m for > 30 seconds', style='List Number')
doc.add_paragraph('4. Alert record is created in database', style='List Number')
doc.add_paragraph('5. Admin dashboard receives real-time notification via Socket.io', style='List Number')
doc.add_paragraph('6. Emergency contacts receive SMS via Twilio API', style='List Number')
doc.add_paragraph('7. Admin reviews alert and initiates response', style='List Number')
doc.add_paragraph('8. Alert is marked as resolved after verification', style='List Number')

doc.add_page_break()

# ==================== 3. CLASS DIAGRAM ====================
doc.add_heading('3. CLASS DIAGRAM', level=1)
doc.add_paragraph('The class diagram shows the object-oriented structure of the SAFORA system:')
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Core Classes:\n').font.bold = True
doc.add_paragraph()

# User Class
p = doc.add_paragraph()
p.add_run('User (Base Class)\n').font.bold = True
doc.add_paragraph('Attributes: id, name, email, phone, password, cnic, gender, role', style='List Bullet')
doc.add_paragraph('Methods: register(), login(), updateProfile(), logout()', style='List Bullet')
doc.add_paragraph()

# Passenger Class
p = doc.add_paragraph()
p.add_run('Passenger extends User\n').font.bold = True
doc.add_paragraph('Attributes: emergencyContacts[], pinkPassVerified, verificationDate', style='List Bullet')
doc.add_paragraph('Methods: enrollPinkPass(), requestRide(), trackRide(), rateDriver(), triggerSOS()', style='List Bullet')
doc.add_paragraph()

# Driver Class
p = doc.add_paragraph()
p.add_run('Driver extends User\n').font.bold = True
doc.add_paragraph('Attributes: licenseNumber, vehicleInfo, rating, totalRides, status, currentLocation, idleTime, onlineTime', style='List Bullet')
doc.add_paragraph('Methods: updateStatus(), updateLocation(), acceptRide(), rejectRide(), startRide(), completeRide()', style='List Bullet')
doc.add_paragraph()

# Admin Class
p = doc.add_paragraph()
p.add_run('Admin extends User\n').font.bold = True
doc.add_paragraph('Attributes: permissions[], department', style='List Bullet')
doc.add_paragraph('Methods: monitorRides(), viewAlerts(), manageUsers(), generateReports(), resolveAlert()', style='List Bullet')
doc.add_paragraph()

# Ride Class
p = doc.add_paragraph()
p.add_run('Ride\n').font.bold = True
doc.add_paragraph('Attributes: id, passenger, driver, pickupLocation, dropoffLocation, plannedRoute, distance, estimatedPrice, actualPrice, status, type, timestamps', style='List Bullet')
doc.add_paragraph('Methods: create(), updateStatus(), calculatePrice(), assignDriver(), trackProgress(), complete()', style='List Bullet')
doc.add_paragraph()

# Alert Class
p = doc.add_paragraph()
p.add_run('Alert\n').font.bold = True
doc.add_paragraph('Attributes: id, ride, type, severity, location, description, status, notificationsSent, resolvedAt', style='List Bullet')
doc.add_paragraph('Methods: create(), notify(), escalate(), resolve(), addNote()', style='List Bullet')

doc.add_page_break()

# ==================== 4. ENTITY RELATIONSHIP DIAGRAM ====================
doc.add_heading('4. ENTITY RELATIONSHIP DIAGRAM', level=1)
doc.add_paragraph('The ER diagram shows the database schema and relationships:')
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Entities and Relationships:\n').font.bold = True
doc.add_paragraph()

# Users Entity
p = doc.add_paragraph()
p.add_run('Users\n').font.bold = True
doc.add_paragraph('Primary Key: _id (ObjectId)', style='List Bullet')
doc.add_paragraph('Attributes: name, email, phone, password (hashed), cnic, gender, role, verified, createdAt', style='List Bullet')
doc.add_paragraph('Relationships: One-to-One with Drivers, One-to-Many with Rides (as passenger)', style='List Bullet')
doc.add_paragraph()

# Drivers Entity
p = doc.add_paragraph()
p.add_run('Drivers\n').font.bold = True
doc.add_paragraph('Primary Key: _id (ObjectId)', style='List Bullet')
doc.add_paragraph('Foreign Key: user (ref: Users)', style='List Bullet')
doc.add_paragraph('Attributes: licenseNumber, vehicleInfo, rating, totalRides, status, currentLocation (GeoJSON), idleTime, onlineTime', style='List Bullet')
doc.add_paragraph('Relationships: One-to-One with Users, One-to-Many with Rides (as driver)', style='List Bullet')
doc.add_paragraph()

# Rides Entity
p = doc.add_paragraph()
p.add_run('Rides\n').font.bold = True
doc.add_paragraph('Primary Key: _id (ObjectId)', style='List Bullet')
doc.add_paragraph('Foreign Keys: passenger (ref: Users), driver (ref: Drivers)', style='List Bullet')
doc.add_paragraph('Attributes: pickupLocation, dropoffLocation, plannedRoute, distance, estimatedPrice, actualPrice, status, type, timestamps', style='List Bullet')
doc.add_paragraph('Relationships: Many-to-One with Users (passenger), Many-to-One with Drivers, One-to-Many with Alerts', style='List Bullet')
doc.add_paragraph()

# Alerts Entity
p = doc.add_paragraph()
p.add_run('Alerts\n').font.bold = True
doc.add_paragraph('Primary Key: _id (ObjectId)', style='List Bullet')
doc.add_paragraph('Foreign Key: ride (ref: Rides)', style='List Bullet')
doc.add_paragraph('Attributes: type, severity, location, description, status, notificationsSent, resolvedAt, notes', style='List Bullet')
doc.add_paragraph('Relationships: Many-to-One with Rides', style='List Bullet')

doc.add_page_break()

# ==================== 5. ACTIVITY DIAGRAMS ====================
doc.add_heading('5. ACTIVITY DIAGRAMS', level=1)

doc.add_heading('5.1 Ride Booking Process', level=2)
doc.add_paragraph('This activity diagram shows the complete ride booking workflow:')
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Process Steps:\n').font.bold = True
doc.add_paragraph('START', style='List Number')
doc.add_paragraph('Passenger opens app and logs in', style='List Number')
doc.add_paragraph('Passenger selects pickup location', style='List Number')
doc.add_paragraph('Passenger selects dropoff location', style='List Number')
doc.add_paragraph('[Decision] Pink Pass ride?', style='List Number')
doc.add_paragraph('  - YES: Filter for verified female drivers', style='List Bullet')
doc.add_paragraph('  - NO: Search all available drivers', style='List Bullet')
doc.add_paragraph('System calculates route and price', style='List Number')
doc.add_paragraph('System matches driver using AI algorithm', style='List Number')
doc.add_paragraph('[Decision] Driver accepts?', style='List Number')
doc.add_paragraph('  - YES: Proceed to next step', style='List Bullet')
doc.add_paragraph('  - NO: Match next available driver', style='List Bullet')
doc.add_paragraph('Passenger receives driver details', style='List Number')
doc.add_paragraph('Driver navigates to pickup', style='List Number')
doc.add_paragraph('Driver arrives and starts ride', style='List Number')
doc.add_paragraph('System monitors GPS during trip', style='List Number')
doc.add_paragraph('Driver arrives at destination', style='List Number')
doc.add_paragraph('Passenger confirms arrival', style='List Number')
doc.add_paragraph('Payment is processed', style='List Number')
doc.add_paragraph('Passenger rates driver', style='List Number')
doc.add_paragraph('END', style='List Number')

doc.add_page_break()

doc.add_heading('5.2 Emergency Alert Process', level=2)
doc.add_paragraph('This activity diagram shows the emergency alert workflow:')
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Process Steps:\n').font.bold = True
doc.add_paragraph('START (Active Ride in Progress)', style='List Number')
doc.add_paragraph('Safety Sentinel monitors GPS location', style='List Number')
doc.add_paragraph('[Decision] Route deviation detected?', style='List Number')
doc.add_paragraph('  - NO: Continue monitoring', style='List Bullet')
doc.add_paragraph('  - YES: Start deviation timer', style='List Bullet')
doc.add_paragraph('[Decision] Deviation exceeds 30 seconds?', style='List Number')
doc.add_paragraph('  - NO: Wait and recheck', style='List Bullet')
doc.add_paragraph('  - YES: Create alert record', style='List Bullet')
doc.add_paragraph('[Parallel Process] Send notifications:', style='List Number')
doc.add_paragraph('  - Send Socket.io notification to admin dashboard', style='List Bullet')
doc.add_paragraph('  - Send SMS to passenger emergency contacts', style='List Bullet')
doc.add_paragraph('Admin receives and reviews alert', style='List Number')
doc.add_paragraph('[Decision] False alarm?', style='List Number')
doc.add_paragraph('  - YES: Mark as resolved (false positive)', style='List Bullet')
doc.add_paragraph('  - NO: Initiate emergency response', style='List Bullet')
doc.add_paragraph('Admin adds resolution notes', style='List Number')
doc.add_paragraph('Alert is closed in system', style='List Number')
doc.add_paragraph('END', style='List Number')

doc.add_page_break()

# ==================== 6. STATE DIAGRAMS ====================
doc.add_heading('6. STATE DIAGRAMS', level=1)

doc.add_heading('6.1 Ride State Diagram', level=2)
doc.add_paragraph('This state diagram shows all possible states of a ride:')
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('States and Transitions:\n').font.bold = True
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('1. REQUESTED\n').font.bold = True
doc.add_paragraph('  → MATCHED (when driver accepts)', style='List Bullet')
doc.add_paragraph('  → CANCELLED (if passenger cancels)', style='List Bullet')
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('2. MATCHED\n').font.bold = True
doc.add_paragraph('  → DRIVER_EN_ROUTE (driver starts navigation)', style='List Bullet')
doc.add_paragraph('  → CANCELLED (if driver/passenger cancels)', style='List Bullet')
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('3. DRIVER_EN_ROUTE\n').font.bold = True
doc.add_paragraph('  → DRIVER_ARRIVED (driver reaches pickup location)', style='List Bullet')
doc.add_paragraph('  → CANCELLED (if passenger no-show)', style='List Bullet')
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('4. DRIVER_ARRIVED\n').font.bold = True
doc.add_paragraph('  → IN_PROGRESS (passenger boards and ride starts)', style='List Bullet')
doc.add_paragraph('  → CANCELLED (if passenger doesn\'t show within timeout)', style='List Bullet')
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('5. IN_PROGRESS\n').font.bold = True
doc.add_paragraph('  → COMPLETED (destination reached)', style='List Bullet')
doc.add_paragraph('  → CANCELLED (emergency cancellation)', style='List Bullet')
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('6. COMPLETED\n').font.bold = True
doc.add_paragraph('  Final state - ride successfully finished', style='List Bullet')
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('7. CANCELLED\n').font.bold = True
doc.add_paragraph('  Final state - ride was cancelled at some stage', style='List Bullet')

doc.add_page_break()

doc.add_heading('6.2 Driver State Diagram', level=2)
doc.add_paragraph('This state diagram shows the operational states of a driver:')
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('States and Transitions:\n').font.bold = True
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('1. OFFLINE\n').font.bold = True
doc.add_paragraph('  → ONLINE (driver activates availability)', style='List Bullet')
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('2. ONLINE\n').font.bold = True
doc.add_paragraph('  → RIDE_ASSIGNED (receives and accepts ride request)', style='List Bullet')
doc.add_paragraph('  → OFFLINE (driver deactivates availability)', style='List Bullet')
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('3. RIDE_ASSIGNED\n').font.bold = True
doc.add_paragraph('  → EN_ROUTE_TO_PICKUP (starts navigation to passenger)', style='List Bullet')
doc.add_paragraph('  → ONLINE (ride cancelled before pickup)', style='List Bullet')
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('4. EN_ROUTE_TO_PICKUP\n').font.bold = True
doc.add_paragraph('  → AT_PICKUP (arrives at passenger location)', style='List Bullet')
doc.add_paragraph('  → ONLINE (ride cancelled)', style='List Bullet')
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('5. AT_PICKUP\n').font.bold = True
doc.add_paragraph('  → ON_TRIP (passenger boards, trip starts)', style='List Bullet')
doc.add_paragraph('  → ONLINE (passenger no-show timeout)', style='List Bullet')
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('6. ON_TRIP\n').font.bold = True
doc.add_paragraph('  → ONLINE (trip completed, passenger dropped off)', style='List Bullet')

doc.add_page_break()

# ==================== 7. COMPONENT DIAGRAM ====================
doc.add_heading('7. COMPONENT DIAGRAM', level=1)
doc.add_paragraph('The component diagram shows the system architecture and component interactions:')
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('System Components:\n').font.bold = True
doc.add_paragraph()

# Mobile Apps
p = doc.add_paragraph()
p.add_run('1. Mobile Applications (React Native)\n').font.bold = True
doc.add_paragraph('Passenger App', style='List Bullet')
doc.add_paragraph('  - Authentication Module', style='List Bullet')
doc.add_paragraph('  - Ride Booking Module', style='List Bullet')
doc.add_paragraph('  - Pink Pass Enrollment Module', style='List Bullet')
doc.add_paragraph('  - Real-time Tracking Module', style='List Bullet')
doc.add_paragraph('  - Rating & Feedback Module', style='List Bullet')
doc.add_paragraph('Driver App', style='List Bullet')
doc.add_paragraph('  - Authentication Module', style='List Bullet')
doc.add_paragraph('  - Ride Management Module', style='List Bullet')
doc.add_paragraph('  - Navigation Module', style='List Bullet')
doc.add_paragraph('  - Earnings Module', style='List Bullet')
doc.add_paragraph()

# Admin Dashboard
p = doc.add_paragraph()
p.add_run('2. Admin Dashboard (React.js)\n').font.bold = True
doc.add_paragraph('  - Monitoring Module', style='List Bullet')
doc.add_paragraph('  - Alert Management Module', style='List Bullet')
doc.add_paragraph('  - User Management Module', style='List Bullet')
doc.add_paragraph('  - Analytics Module', style='List Bullet')
doc.add_paragraph('  - Heatmap Visualization Module', style='List Bullet')
doc.add_paragraph()

# Backend Server
p = doc.add_paragraph()
p.add_run('3. Backend Server (Node.js/Express)\n').font.bold = True
doc.add_paragraph('  - API Gateway', style='List Bullet')
doc.add_paragraph('  - Authentication Service (JWT)', style='List Bullet')
doc.add_paragraph('  - User Management Service', style='List Bullet')
doc.add_paragraph('  - Ride Management Service', style='List Bullet')
doc.add_paragraph('  - Safety Sentinel Service', style='List Bullet')
doc.add_paragraph('  - Notification Service (Socket.io)', style='List Bullet')
doc.add_paragraph('  - SMS Service (Twilio)', style='List Bullet')
doc.add_paragraph()

# AI Microservice
p = doc.add_paragraph()
p.add_run('4. AI Microservice (Python/Flask)\n').font.bold = True
doc.add_paragraph('  - Pink Pass Verification Service', style='List Bullet')
doc.add_paragraph('  - Price Prediction Service', style='List Bullet')
doc.add_paragraph('  - Driver Matching Service', style='List Bullet')
doc.add_paragraph('  - Sentiment Analysis Service', style='List Bullet')
doc.add_paragraph()

# Database
p = doc.add_paragraph()
p.add_run('5. Database Layer (MongoDB)\n').font.bold = True
doc.add_paragraph('  - Users Collection', style='List Bullet')
doc.add_paragraph('  - Drivers Collection', style='List Bullet')
doc.add_paragraph('  - Rides Collection', style='List Bullet')
doc.add_paragraph('  - Alerts Collection', style='List Bullet')
doc.add_paragraph()

# External APIs
p = doc.add_paragraph()
p.add_run('6. External APIs\n').font.bold = True
doc.add_paragraph('  - Google Maps API (routing, geocoding)', style='List Bullet')
doc.add_paragraph('  - Twilio API (SMS notifications)', style='List Bullet')

doc.add_page_break()

# ==================== 8. DEPLOYMENT DIAGRAM ====================
doc.add_heading('8. DEPLOYMENT DIAGRAM', level=1)
doc.add_paragraph('The deployment diagram shows the physical deployment architecture:')
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Deployment Architecture:\n').font.bold = True
doc.add_paragraph()

# Client Devices
p = doc.add_paragraph()
p.add_run('1. Client Devices\n').font.bold = True
doc.add_paragraph('Mobile Devices (iOS/Android)', style='List Bullet')
doc.add_paragraph('  - Passenger App', style='List Bullet')
doc.add_paragraph('  - Driver App', style='List Bullet')
doc.add_paragraph('Web Browsers', style='List Bullet')
doc.add_paragraph('  - Admin Dashboard', style='List Bullet')
doc.add_paragraph()

# Application Server
p = doc.add_paragraph()
p.add_run('2. Application Server (Cloud - AWS/DigitalOcean)\n').font.bold = True
doc.add_paragraph('Node.js Backend Server', style='List Bullet')
doc.add_paragraph('  - Express.js API', style='List Bullet')
doc.add_paragraph('  - Socket.io WebSocket Server', style='List Bullet')
doc.add_paragraph('  - Port: 5000', style='List Bullet')
doc.add_paragraph('Python AI Service', style='List Bullet')
doc.add_paragraph('  - Flask Application', style='List Bullet')
doc.add_paragraph('  - ML Models (TensorFlow/Keras)', style='List Bullet')
doc.add_paragraph('  - Port: 5001', style='List Bullet')
doc.add_paragraph()

# Database Server
p = doc.add_paragraph()
p.add_run('3. Database Server (Cloud/On-Premise)\n').font.bold = True
doc.add_paragraph('MongoDB Instance', style='List Bullet')
doc.add_paragraph('  - Port: 27017', style='List Bullet')
doc.add_paragraph('  - Collections: Users, Drivers, Rides, Alerts', style='List Bullet')
doc.add_paragraph('  - Replication: Enabled for high availability', style='List Bullet')
doc.add_paragraph()

# External Services
p = doc.add_paragraph()
p.add_run('4. External Services\n').font.bold = True
doc.add_paragraph('Google Maps Platform', style='List Bullet')
doc.add_paragraph('  - Maps API', style='List Bullet')
doc.add_paragraph('  - Directions API', style='List Bullet')
doc.add_paragraph('  - Geocoding API', style='List Bullet')
doc.add_paragraph('Twilio Service', style='List Bullet')
doc.add_paragraph('  - SMS API', style='List Bullet')
doc.add_paragraph()

# Communication Protocols
p = doc.add_paragraph()
p.add_run('Communication Protocols:\n').font.bold = True
doc.add_paragraph('HTTPS/TLS 1.3 - Secure API communication', style='List Bullet')
doc.add_paragraph('WebSocket (Socket.io) - Real-time updates', style='List Bullet')
doc.add_paragraph('MongoDB Protocol - Database connections', style='List Bullet')
doc.add_paragraph('REST API - External service integration', style='List Bullet')

doc.add_page_break()

# ==================== DOCUMENT APPROVAL ====================
doc.add_heading('Document Approval', level=1)
doc.add_paragraph('This SRS Phase 2 document has been reviewed and approved by:')
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Project Advisor:\n').font.bold = True
p.add_run('Tanveer Ahmed\n')
p.add_run('Date: _______________')
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Team Members:\n').font.bold = True
p.add_run('Mirza Umer Ikram (S3F22UBSCS081)\n')
p.add_run('Date: _______________')
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Ruhma Bilal (S3F22UBSCS088)\n')
p.add_run('Date: _______________')

# Save document
output_path = 'SAFORA_SRS2_System_Design_Diagrams.docx'
doc.save(output_path)
print(f"✅ SRS2 document generated: {output_path}")
print("✅ Document includes:")
print("   - Use Case Diagrams (Passenger, Driver, Admin)")
print("   - Sequence Diagrams (Registration, Ride Request, Pink Pass, Safety Alert)")
print("   - Class Diagram")
print("   - Entity Relationship Diagram")
print("   - Activity Diagrams (Ride Booking, Emergency Alert)")
print("   - State Diagrams (Ride, Driver)")
print("   - Component Diagram")
print("   - Deployment Diagram")
