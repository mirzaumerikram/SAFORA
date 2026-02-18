from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
import sys

# Fix console encoding for Windows
if sys.platform == 'win32':
    import codecs
    sys.stdout = codecs.getwriter('utf-8')(sys.stdout.buffer, 'strict')

def add_gantt_chart_to_section6():
    """Add comprehensive Gantt chart and updated content to Section 6"""
    
    # Load the document
    doc = Document('SAFORA_SRS1_IEEE_Format_Complete.docx')
    
    # Find Section 6
    section_6_para = None
    section_6_index = None
    for i, para in enumerate(doc.paragraphs):
        if '6.' in para.text[:5] and 'Revised' in para.text:
            section_6_para = para
            section_6_index = i
            break
    
    if section_6_index is None:
        print("Section 6 not found!")
        return
    
    print(f"Found Section 6 at paragraph index: {section_6_index}")
    
    # Find the next section (Section 7) to know where to stop
    section_7_index = None
    for i in range(section_6_index + 1, len(doc.paragraphs)):
        if doc.paragraphs[i].text.strip().startswith('7.'):
            section_7_index = i
            break
    
    print(f"Found Section 7 at paragraph index: {section_7_index}")
    
    # Delete existing content between Section 6 and Section 7
    if section_7_index:
        for _ in range(section_7_index - section_6_index - 1):
            p = doc.paragraphs[section_6_index + 1]
            p._element.getparent().remove(p._element)
    
    # Now insert new content after Section 6 heading
    # We'll collect all paragraphs first, then insert them
    
    new_paragraphs = []
    
    # 6.1 Project Overview
    new_paragraphs.append(("heading2", "6.1 Project Overview"))
    new_paragraphs.append(("normal", """The SAFORA (Safe and Affordable Ride-Sharing for All) project follows an Agile development methodology with 2-week sprint cycles. The project is divided into five major phases spanning from September 2024 to May 2025, totaling 9 months of development. The team consists of 3 core members working collaboratively on backend, frontend, and AI/ML components."""))
    
    # 6.2 Current Project Status
    new_paragraphs.append(("heading2", "6.2 Current Project Status (As of January 2025)"))
    new_paragraphs.append(("normal", """As of January 2025, the project has successfully completed Phase 1 (Requirements & Planning) and is currently in the final stages of Phase 2 (Backend Development). The overall project completion is approximately 40%, with the following phase-wise breakdown:

• Phase 1: Requirements & Planning - 100% Complete
• Phase 2: Backend Development - 75% Complete (In Progress)
• Phase 3: AI/ML Integration - 0% (Scheduled to begin February 2025)
• Phase 4: Mobile Apps - 0% (Scheduled for March-April 2025)
• Phase 5: Testing & Deployment - 0% (Scheduled for May 2025)"""))
    
    # 6.3 Project Timeline - Gantt Chart
    new_paragraphs.append(("heading2", "6.3 Project Timeline - Gantt Chart"))
    new_paragraphs.append(("normal", """Figure 6.1 below presents the detailed Gantt chart illustrating the project timeline, task dependencies, and current progress across all five phases. The chart uses color coding to indicate task status: green for completed tasks, orange for in-progress tasks, and gray for tasks not yet started."""))
    new_paragraphs.append(("image", r"C:\Users\mzees\.gemini\antigravity\brain\6e3684ad-553b-4832-9cc1-ea3bd890ddb4\gantt_safora_project_1769696917006.png"))
    new_paragraphs.append(("caption", "Figure 6.1: SAFORA Project Gantt Chart (September 2024 - May 2025)"))
    
    # 6.4 Phase-wise Detailed Progress
    new_paragraphs.append(("heading2", "6.4 Phase-wise Detailed Progress"))
    
    new_paragraphs.append(("normal", """6.4.1 Phase 1: Requirements & Planning (Completed - 100%)

Duration: September - October 2024 (Weeks 1-8)
Status: Completed

Completed Deliverables:
• Project Proposal with comprehensive problem analysis
• Requirements Specification Document (SRS1) following IEEE 830 standards
• System Design Document (SRS2) with UML diagrams
• Technology stack selection and justification
• Risk assessment and mitigation strategies

All deliverables for Phase 1 were completed on schedule, providing a solid foundation for development phases."""))
    
    new_paragraphs.append(("normal", """6.4.2 Phase 2: Backend Development (In Progress - 75%)

Duration: November - December 2024 (Weeks 9-16)
Status: In Progress (Currently at Week 15)

Completed Tasks:
• Node.js/Express server architecture setup with RESTful API design
• MongoDB database schema implementation with proper indexing
• JWT-based authentication system with role-based access control
• User management APIs for Passenger, Driver, and Admin roles
• Ride request and matching algorithm implementation
• Socket.io integration for real-time communication

In Progress:
• Safety Features Backend (GPS tracking, Safety Sentinel, SOS alerts)
• Integration with Twilio API for SMS notifications
• Google Maps API integration for routing and geocoding

The backend core is on track with minor adjustments to the safety features timeline due to API integration complexity."""))
    
    new_paragraphs.append(("normal", """6.4.3 Phase 3: AI/ML Integration (Scheduled - 0%)

Duration: January - February 2025 (Weeks 17-24)
Status: Scheduled to begin February 1, 2025

Planned Tasks:
• Python/Flask microservice architecture setup
• Sentiment analysis model for ride feedback (using TextBlob)
• Route optimization AI using pathfinding algorithms
• Safety Sentinel system with anomaly detection
• Machine learning model training and validation
• AI service integration with backend APIs

Dependencies: Completion of Phase 2 backend safety features."""))
    
    new_paragraphs.append(("normal", """6.4.4 Phase 4: Mobile Applications (Scheduled - 0%)

Duration: March - April 2025 (Weeks 25-32)
Status: Scheduled to begin March 1, 2025

Planned Tasks:
• React Native setup for cross-platform mobile development
• Passenger App: Ride booking, tracking, payment interface
• Driver App: Trip management, navigation, earnings dashboard
• Real-time features integration (GPS tracking, chat, notifications)
• UI/UX refinement and accessibility compliance
• Mobile app testing on iOS and Android devices

Dependencies: Completion of Phase 2 and Phase 3."""))
    
    new_paragraphs.append(("normal", """6.4.5 Phase 5: Testing & Deployment (Scheduled - 0%)

Duration: May 2025 (Weeks 33-36)
Status: Scheduled to begin May 1, 2025

Planned Tasks:
• System integration testing across all components
• Performance and load testing
• Security testing and vulnerability assessment
• User acceptance testing (UAT) with pilot users
• Production deployment and monitoring setup
• Final documentation and project handover
• Viva and final presentation preparation

Dependencies: Completion of all development phases."""))
    
    # 6.5 Updated Project Plan & Deviations
    new_paragraphs.append(("heading2", "6.5 Updated Project Plan & Deviations from Original Proposal"))
    new_paragraphs.append(("normal", """The project is largely on track with the original proposal timeline submitted in September 2024. However, the following minor adjustments have been made:

Modifications to Original Plan:
1. Backend Safety Features: Extended by 1 week due to complexity in integrating Twilio SMS and Google Maps APIs for real-time safety alerts.

2. AI Service Architecture: Decision to use Python/Flask microservice instead of integrating AI directly into Node.js backend for better separation of concerns and scalability.

3. Technology Stack Addition: Socket.io added for enhanced real-time communication capabilities not initially specified in the proposal.

Risk Mitigation Actions:
• Regular bi-weekly sprint reviews to identify and address blockers early
• Parallel development tracks to minimize dependencies
• Comprehensive documentation maintained throughout development
• Continuous integration for early detection of integration issues

Planned Buffer Time:
• 2-week buffer built into Phase 5 (Testing) for unexpected delays
• Flexible sprint allocation to accommodate learning curve for new technologies

Despite these minor adjustments, all major milestones remain achievable within the May 2025 deadline."""))
    
    # 6.6 Resource Allocation
    new_paragraphs.append(("heading2", "6.6 Resource Allocation"))
    new_paragraphs.append(("normal", """Team Composition: 3 members (Full-Stack Developers with specialized roles)

Role Distribution:
• Backend Lead: Node.js/Express API development, database design, authentication
• Frontend Lead: React Native mobile apps, UI/UX design, user experience
• AI/ML Lead: Python/Flask AI service, machine learning models, data analysis

Shared Responsibilities:
• All team members contribute to code reviews and testing
• Weekly standup meetings for progress synchronization
• Joint responsibility for integration testing and deployment

Development Tools & Infrastructure:
• Version Control: Git/GitHub with feature branch workflow
• Project Management: Jira for sprint planning and task tracking
• Communication: Slack for team coordination, Zoom for remote meetings
• Development Environment: VS Code, Postman, MongoDB Compass
• Cloud Services: To be determined (AWS/Azure/Google Cloud for deployment)"""))
    
    # Insert all new paragraphs after Section 6 heading
    insert_position = section_6_para._element
    parent = insert_position.getparent()
    
    for para_type, content in new_paragraphs:
        if para_type == "heading2":
            p = doc.add_paragraph(content)
            p.style = 'Heading 2'
            parent.insert(parent.index(insert_position) + 1, p._element)
            insert_position = p._element
        elif para_type == "normal":
            p = doc.add_paragraph(content)
            parent.insert(parent.index(insert_position) + 1, p._element)
            insert_position = p._element
        elif para_type == "image":
            if os.path.exists(content):
                p = doc.add_paragraph()
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = p.add_run()
                run.add_picture(content, width=Inches(6.5))
                parent.insert(parent.index(insert_position) + 1, p._element)
                insert_position = p._element
            else:
                print(f"Warning: Image not found at {content}")
        elif para_type == "caption":
            p = doc.add_paragraph(content)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = p.runs[0]
            run.font.italic = True
            run.font.size = Pt(10)
            parent.insert(parent.index(insert_position) + 1, p._element)
            insert_position = p._element
    
    # Save the updated document
    output_path = 'SAFORA_SRS1_IEEE_Format_Complete_Updated.docx'
    doc.save(output_path)
    print(f"\n[SUCCESS] Document updated and saved as: {output_path}")
    print("\n[ADDED] Section 6 Content:")
    print("  - 6.1 Project Overview")
    print("  - 6.2 Current Project Status")
    print("  - 6.3 Project Timeline - Gantt Chart (with visual diagram)")
    print("  - 6.4 Phase-wise Detailed Progress (5 subsections)")
    print("  - 6.5 Updated Project Plan & Deviations")
    print("  - 6.6 Resource Allocation")
    print(f"\n[INFO] Total new content added: {len(new_paragraphs)} elements")

if __name__ == "__main__":
    add_gantt_chart_to_section6()
