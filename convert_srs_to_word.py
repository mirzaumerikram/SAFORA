from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def add_table_border(table):
    """Add borders to table"""
    tbl = table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)

# Create a new Document
doc = Document()

# Set up styles
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# Title Page
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run('SAFORA (Smart Ride)\n')
title_run.font.size = Pt(18)
title_run.font.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.add_run('Software Requirements Specification - Phase 1')
subtitle_run.font.size = Pt(16)
subtitle_run.font.bold = True

doc.add_paragraph()
doc.add_paragraph()

# University info
uni = doc.add_paragraph()
uni.alignment = WD_ALIGN_PARAGRAPH.CENTER
uni.add_run('University of Central Punjab\n').font.bold = True
uni.add_run('Faculty of Information Technology\n')
uni.add_run('Department of Computer Science')

doc.add_paragraph()
doc.add_paragraph()

# Project details
details = doc.add_paragraph()
details.alignment = WD_ALIGN_PARAGRAPH.CENTER
details.add_run('Project Title: ').font.bold = True
details.add_run('SAFORA - Smart Ride\n')
details.add_run('Project Advisor: ').font.bold = True
details.add_run('Tanveer Ahmed\n')
details.add_run('Group ID: ').font.bold = True
details.add_run('11\n\n')
details.add_run('Team Members:\n').font.bold = True
details.add_run('Mirza Umer Ikram (S3F22UBSCS081)\n')
details.add_run('Ruhma Bilal (S3F22UBSCS088)\n\n')
details.add_run('Date: January 19, 2026')

# Page break
doc.add_page_break()

# Read the markdown content
with open('SAFORA_SRS1_IEEE_Format.md', 'r', encoding='utf-8') as f:
    content = f.read()

# Split content into lines
lines = content.split('\n')
skip_until_intro = True
table_lines = []
in_table = False

i = 0
while i < len(lines):
    line = lines[i].strip()
    
    # Skip title page content
    if '# 1. INTRODUCTION' in line:
        skip_until_intro = False
    
    if skip_until_intro:
        i += 1
        continue
    
    # Detect table start
    if line.startswith('|') and i + 1 < len(lines) and lines[i + 1].strip().startswith('|'):
        in_table = True
        table_lines = [line]
        i += 1
        
        # Collect all table lines
        while i < len(lines) and lines[i].strip().startswith('|'):
            table_lines.append(lines[i].strip())
            i += 1
        
        # Parse and create table
        if len(table_lines) >= 2:
            # Parse header
            header = [cell.strip() for cell in table_lines[0].split('|')[1:-1]]
            
            # Skip separator line
            rows = []
            for table_line in table_lines[2:]:
                row = [cell.strip() for cell in table_line.split('|')[1:-1]]
                if row and any(row):  # Skip empty rows
                    rows.append(row)
            
            # Create table
            if rows:
                table = doc.add_table(rows=len(rows) + 1, cols=len(header))
                add_table_border(table)
                
                # Add header
                for j, cell_text in enumerate(header):
                    cell = table.rows[0].cells[j]
                    cell.text = cell_text
                    cell.paragraphs[0].runs[0].font.bold = True
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Add rows
                for row_idx, row_data in enumerate(rows):
                    for col_idx, cell_text in enumerate(row_data):
                        table.rows[row_idx + 1].cells[col_idx].text = cell_text
                
                doc.add_paragraph()  # Add spacing after table
        
        in_table = False
        table_lines = []
        continue
    
    if not line:
        i += 1
        continue
    
    # Heading 1 (# )
    if line.startswith('# ') and not line.startswith('## '):
        heading = doc.add_heading(line[2:], level=1)
        heading.runs[0].font.size = Pt(16)
        heading.runs[0].font.bold = True
    
    # Heading 2 (## )
    elif line.startswith('## ') and not line.startswith('### '):
        heading = doc.add_heading(line[3:], level=2)
        heading.runs[0].font.size = Pt(14)
        heading.runs[0].font.bold = True
    
    # Heading 3 (### )
    elif line.startswith('### '):
        heading = doc.add_heading(line[4:], level=3)
        heading.runs[0].font.size = Pt(13)
        heading.runs[0].font.bold = True
    
    # Bold text (**text**)
    elif line.startswith('**') and line.endswith('**'):
        p = doc.add_paragraph()
        p.add_run(line[2:-2]).font.bold = True
    
    # Horizontal rule
    elif line.startswith('---'):
        doc.add_paragraph('_' * 50)
    
    # List items (-)
    elif line.startswith('- '):
        p = doc.add_paragraph(line[2:], style='List Bullet')
    
    # Numbered lists
    elif len(line) > 0 and line[0].isdigit() and '. ' in line:
        p = doc.add_paragraph(line.split('. ', 1)[1], style='List Number')
    
    # Regular paragraph
    else:
        if line and not line.startswith('#'):
            # Handle bold text within paragraphs
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    p.add_run(part[2:-2]).font.bold = True
                elif part:
                    p.add_run(part)
    
    i += 1

# Save the document
output_path = 'SAFORA_SRS1_IEEE_Format.docx'
doc.save(output_path)
print(f"Document saved as: {output_path}")
print("Tables have been properly formatted!")
