import re
import os
from docx import Document
from docx.shared import Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_border(cell):
    """Adds borders to a table cell"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for border in ['top', 'left', 'bottom', 'right']:
        edge = OxmlElement(f'w:{border}')
        edge.set(qn('w:val'), 'single')
        edge.set(qn('w:sz'), '4')
        edge.set(qn('w:space'), '0')
        edge.set(qn('w:color'), '000000')
        tcPr.append(edge)

def add_use_case_to_docx(doc, uc):
    """Adds a single use case to the DOCX with the requested format"""
    doc.add_heading(f"3.{uc['Index']} {uc['Name']}", level=2)
    
    # 1. Base Info Table
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.columns[0].width = Cm(4)
    table.columns[1].width = Cm(13)

    fields = [
        ("Identifier", uc['Identifier']),
        ("Purpose", uc['Purpose']),
        ("Priority", uc['Priority']),
        ("Pre-conditions", uc['Pre-conditions']),
        ("Post-conditions", uc['Post-conditions'])
    ]

    for i, (label, val) in enumerate(fields):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        row.cells[1].text = val
    
    doc.add_paragraph()

    # 2. Typical Course of Action
    if uc['Typical']:
        p = doc.add_paragraph()
        p.add_run("Typical Course of Action").bold = True
        
        typ_table = doc.add_table(rows=len(uc['Typical']) + 1, cols=3)
        typ_table.style = 'Table Grid'
        typ_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        typ_table.autofit = False
        typ_table.columns[0].width = Cm(1)
        typ_table.columns[1].width = Cm(8)
        typ_table.columns[2].width = Cm(8)

        # Headers
        hdr = typ_table.rows[0].cells
        hdr[0].text = "S#"
        hdr[1].text = "Actor Action"
        hdr[2].text = "System Response"
        for cell in hdr:
            cell.paragraphs[0].runs[0].font.bold = True

        for i, step in enumerate(uc['Typical']):
            row = typ_table.rows[i+1].cells
            row[0].text = step[0]
            row[1].text = step[1]
            row[2].text = step[2]
            
        doc.add_paragraph()

    # 3. Alternate Course of Action
    if uc['Alternate']:
        p = doc.add_paragraph()
        p.add_run("Alternate Course of Action").bold = True
        
        alt_table = doc.add_table(rows=len(uc['Alternate']) + 1, cols=3)
        alt_table.style = 'Table Grid'
        alt_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        alt_table.autofit = False
        alt_table.columns[0].width = Cm(1)
        alt_table.columns[1].width = Cm(8)
        alt_table.columns[2].width = Cm(8)

        # Headers
        hdr = alt_table.rows[0].cells
        hdr[0].text = "S#"
        hdr[1].text = "Actor Action"
        hdr[2].text = "System Response"
        for cell in hdr:
            cell.paragraphs[0].runs[0].font.bold = True

        for i, step in enumerate(uc['Alternate']):
            row = alt_table.rows[i+1].cells
            row[0].text = step[0]
            row[1].text = step[1]
            row[2].text = step[2]

    doc.add_page_break()

def parse_md(file_path):
    with open(file_path, 'r', encoding='utf-8') as f:
        content = f.read()

    sections = re.split(r'### \d+\.\d+ ', content)[1:]
    use_cases = []

    for i, section in enumerate(sections):
        lines = section.split('\n')
        name = lines[0].strip()
        
        uc = {'Index': i+1, 'Name': name, 'Typical': [], 'Alternate': []}

        for field in ['Identifier', 'Purpose', 'Priority', 'Pre-conditions', 'Post-conditions']:
            match = re.search(fr'\*\*{field}\*\* \| (.*?)(?:\s*\|)?$', section, re.MULTILINE)
            uc[field] = match.group(1).strip().replace('|', '').strip() if match else ''

        typ_match = re.search(r'\*\*Typical Course of Action\*\*\n\| S# \| Actor Action \| System Response \|\n\| :--- \| :--- \| :--- \|\n(.*?)(?=\n\n\*\*Alt|\n\n---|$)', section, re.DOTALL)
        if typ_match:
            for line in typ_match.group(1).strip().split('\n'):
                parts = [p.strip() for p in line.split('|')[1:-1]]
                if len(parts) == 3: uc['Typical'].append(parts)

        alt_match = re.search(r'\*\*Alternate Course of Action\*\*\n\| S# \| Actor Action \| System Response \|\n\| :--- \| :--- \| :--- \|\n(.*?)(?=\n\n---|$)', section, re.DOTALL)
        if alt_match:
            for line in alt_match.group(1).strip().split('\n'):
                parts = [p.strip() for p in line.split('|')[1:-1]]
                if len(parts) == 3: uc['Alternate'].append(parts)

        use_cases.append(uc)
    return use_cases

if __name__ == "__main__":
    md_file = "SAFORA_30_Use_Cases.md"
    docx_file = "SAFORA_30_Use_Cases.docx"
    
    if os.path.exists(md_file):
        data = parse_md(md_file)
        doc = Document()
        doc.add_heading('SAFORA - 30 Detailed Use Cases', 0)
        
        for uc in data:
            add_use_case_to_docx(doc, uc)
            
        doc.save(docx_file)
        print(f"Successfully generated {docx_file}")
    else:
        print("Source MD file not found.")
