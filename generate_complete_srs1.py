"""
Complete SRS1 Document Generator for SAFORA Project (v3)
Generates IEEE-formatted Word document with:
- Dynamic Table of Contents
- Headers (on subsequent pages)
- Footers with Page Numbers
- Professional Styling
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_border(cell, **kwargs):
    """Set cell borders"""
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_element = OxmlElement(f'w:{edge}')
        edge_element.set(qn('w:val'), 'single')
        edge_element.set(qn('w:sz'), '4')
        edge_element.set(qn('w:space'), '0')
        edge_element.set(qn('w:color'), '000000')
        tcBorders.append(edge_element)
    
    tcPr.append(tcBorders)

def add_table_with_data(doc, headers, rows, col_widths=None):
    """Create a formatted table"""
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    if col_widths:
        for i, width in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Cm(width)
    
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_border(cell)
    
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_data in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = str(cell_data)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_border(cell)
    
    return table

def create_element(name):
    return OxmlElement(name)

def create_attribute(element, name, value):
    element.set(qn(name), value)

def add_page_number(paragraph):
    """Add dynamic page number"""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Add simple text
    run = paragraph.add_run()
    
    # Start Field
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')
    
    # Field Instr
    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"
    
    # Separator
    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'separate')
    
    # End Field
    fldChar3 = create_element('w:fldChar')
    create_attribute(fldChar3, 'w:fldCharType', 'end')
    
    run._element.append(fldChar1)
    run._element.append(instrText)
    run._element.append(fldChar2)
    run._element.append(fldChar3)

def add_toc(doc):
    """Insert native TOC field"""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')
    
    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    
    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'separate')
    
    fldChar3 = create_element('w:fldChar')
    create_attribute(fldChar3, 'w:fldCharType', 'end')
    
    run._element.append(fldChar1)
    run._element.append(instrText)
    run._element.append(fldChar2)
    run._element.append(fldChar3)

def set_update_fields_on_open(doc):
    """Force TOC update on open"""
    settings = doc.settings.element
    element_proxy = create_element('w:updateFields')
    create_attribute(element_proxy, 'w:val', 'true')
    settings.append(element_proxy)

def add_use_case_table(doc, uc_id, name, actor, description, preconditions, main_flow, postconditions):
    """Create a detailed use case table"""
    table = doc.add_table(rows=7, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    # Set column widths
    table.columns[0].width = Cm(4)
    table.columns[1].width = Cm(12)
    
    # Data rows
    data = [
        ("Use Case ID", uc_id),
        ("Use Case Name", name),
        ("Primary Actor", actor),
        ("Description", description),
        ("Preconditions", preconditions),
        ("Main Flow", main_flow),
        ("Postconditions", postconditions)
    ]
    
    for i, (key, value) in enumerate(data):
        row = table.rows[i]
        
        # Key cell styling
        key_cell = row.cells[0]
        key_cell.text = key
        key_cell.paragraphs[0].runs[0].font.bold = True
        set_cell_border(key_cell)
        
        # Value cell styling
        val_cell = row.cells[1]
        val_cell.text = value
        set_cell_border(val_cell)
        
    doc.add_paragraph()  # Spacing after table

def add_functional_requirement(doc, fr_id, name, description, input=None, processing=None, output=None, priority="High"):
    """Adds a detailed functional requirement in the specific format found in the MD file"""
    p = doc.add_paragraph()
    p.add_run(f'{fr_id}: {name}').font.bold = True
    
    if description:
        dp = doc.add_paragraph(style='List Bullet')
        dp.add_run('Description: ').font.bold = True
        dp.add_run(description)
    
    if input:
        ip = doc.add_paragraph(style='List Bullet')
        ip.add_run('Input: ').font.bold = True
        ip.add_run(input)
        
    if processing:
        pp = doc.add_paragraph(style='List Bullet')
        pp.add_run('Processing: ').font.bold = True
        pp.add_run(processing)
        
    if output:
        op = doc.add_paragraph(style='List Bullet')
        op.add_run('Output: ').font.bold = True
        op.add_run(output)
        
    if priority:
        prp = doc.add_paragraph(style='List Bullet')
        prp.add_run('Priority: ').font.bold = True
        prp.add_run(priority)
    
    doc.add_paragraph() # Spacing

# ==================== MAIN SCRIPT ====================
doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# ==================== HEADERS & FOOTERS ====================
# Configure to have a different first page (Title page has no H/F)
section = doc.sections[0]
section.different_first_page_header_footer = True

# Header for subsequent pages
header = section.header
header_para = header.paragraphs[0]
header_para.text = "SAFORA (Smart Ride) - SRS Phase 1"
header_para.style = doc.styles['Header']
header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# Footer for subsequent pages
footer = section.footer
footer_para = footer.paragraphs[0]
footer_para.text = "Group ID: 11\tPage "
footer_para.style = doc.styles['Footer']

# Add page number field
add_page_number(footer_para)

# Adjust tab stops for footer to align text left and page number right
tab_stops = footer_para.paragraph_format.tab_stops
tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)

# ==================== TITLE PAGE ====================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run('SAFORA (Smart Ride)\n')
title_run.font.size = Pt(20)
title_run.font.bold = True
title_run.font.name = 'Times New Roman'

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.add_run('Software Requirements Specification\nPhase 1')
subtitle_run.font.size = Pt(18)
subtitle_run.font.bold = True
subtitle_run.font.name = 'Times New Roman'

doc.add_paragraph()
doc.add_paragraph()

uni = doc.add_paragraph()
uni.alignment = WD_ALIGN_PARAGRAPH.CENTER
uni_run = uni.add_run('University of Central Punjab\n')
uni_run.font.bold = True
uni_run.font.size = Pt(14)
uni.add_run('Faculty of Information Technology\n').font.size = Pt(12)
uni.add_run('Department of Computer Science').font.size = Pt(12)

doc.add_paragraph()
doc.add_paragraph()

details = doc.add_paragraph()
details.alignment = WD_ALIGN_PARAGRAPH.CENTER
details.add_run('Project Title: ').font.bold = True
details.add_run('SAFORA - Smart Ride\n\n')
details.add_run('Project Advisor: ').font.bold = True
details.add_run('Tanveer Ahmed\n\n')
details.add_run('Group ID: ').font.bold = True
details.add_run('11\n\n')

team = details.add_run('Team Members:\n')
team.font.bold = True
details.add_run('Mirza Umer Ikram (S3F22UBSCS081)\n')
details.add_run('Ruhma Bilal (S3F22UBSCS088)\n\n')

date_run = details.add_run('Date: January 19, 2026')
date_run.font.bold = True

doc.add_page_break()

# ==================== DOCUMENT REVISION HISTORY ====================
doc.add_heading('Document Revision History', level=1)

revision_headers = ['Version', 'Date', 'Description', 'Author(s)']
revision_data = [
    ['1.0', 'January 15, 2026', 'Initial Draft', 'Mirza Umer Ikram, Ruhma Bilal'],
    ['1.1', 'January 19, 2026', 'IEEE Format Update', 'Mirza Umer Ikram, Ruhma Bilal']
]
add_table_with_data(doc, revision_headers, revision_data, [2, 3, 6, 5])
doc.add_paragraph()

# ==================== TABLE OF CONTENTS ====================
doc.add_heading('Table of Contents', level=1)
doc.add_paragraph('Note: Right-click and select "Update Field" to refresh the Table of Contents if needed.')

# Insert native Word TOC field
add_toc(doc)

doc.add_page_break()

# ==================== 1. INTRODUCTION ====================
doc.add_heading('1. INTRODUCTION', level=1)

doc.add_heading('1.1 Purpose', level=2)
doc.add_paragraph(
    'This Software Requirements Specification (SRS) document provides a complete description of the SAFORA '
    '(Smart Ride) system - an AI-powered ride-hailing platform implementing proactive safety measures. '
    'This document is intended for:'
)
doc.add_paragraph('Development team members', style='List Bullet')
doc.add_paragraph('Project stakeholders and supervisors', style='List Bullet')
doc.add_paragraph('Quality assurance personnel', style='List Bullet')
doc.add_paragraph('Future maintenance teams', style='List Bullet')
doc.add_paragraph()
doc.add_paragraph(    'The document follows IEEE Std 830-1998 guidelines [1] for software requirements specifications, '
    'incorporating best practices in modern software engineering [3].')

doc.add_heading('1.2 Scope', level=2)
p = doc.add_paragraph()
p.add_run('Product Name: ').font.bold = True
p.add_run('SAFORA (Smart Ride)')
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Product Description: ').font.bold = True
p.add_run(
    'SAFORA is an intelligent ride-hailing platform that addresses critical safety and efficiency gaps in '
    'existing transportation services through AI-driven proactive safety mechanisms.'
)

doc.add_heading('1.2.1 Product Features', level=3)
doc.add_paragraph('The system implements three core innovations:')

p = doc.add_paragraph()
p.add_run('1. Pink Pass (Biometric Women-Only Mode)').font.bold = True
doc.add_paragraph('Two-layer verification using CNIC and liveness detection', style='List Bullet')
doc.add_paragraph('Haar Cascade for face detection', style='List Bullet')
doc.add_paragraph('CNN-based blink verification for anti-spoofing', style='List Bullet')
doc.add_paragraph('Verification completion in <15 seconds', style='List Bullet')

p = doc.add_paragraph()
p.add_run('2. Safety Sentinel (Automated Route Monitoring)').font.bold = True
doc.add_paragraph('Real-time GPS tracking and route deviation detection', style='List Bullet')
doc.add_paragraph('Point-to-line distance calculation using Turf.js', style='List Bullet')
doc.add_paragraph('Automated alerts when deviation >500m for >30 seconds', style='List Bullet')
doc.add_paragraph('Dual notification system (admin dashboard + SMS to emergency contacts)', style='List Bullet')

p = doc.add_paragraph()
p.add_run('3. Smart Driver Selection').font.bold = True
doc.add_paragraph('AI-driven matching algorithm: Score = 0.50×Distance + 0.30×Rating + 0.20×Fairness', style='List Bullet')
doc.add_paragraph('Fairness metric: IdleTime/TotalOnlineTime (last 4 hours)', style='List Bullet')
doc.add_paragraph('Reduces driver monopolization and wait times', style='List Bullet')

doc.add_heading('1.2.2 Benefits', level=3)
p = doc.add_paragraph()
p.add_run('For Passengers: ').font.bold = True
p.add_run('Enhanced safety through proactive monitoring, verified women-only rides, faster pickup times')

p = doc.add_paragraph()
p.add_run('For Drivers: ').font.bold = True
p.add_run('Fair ride distribution, reduced idle time, transparent rating system')

p = doc.add_paragraph()
p.add_run('For Administrators: ').font.bold = True
p.add_run('Real-time monitoring, automated alerts, data-driven insights')

doc.add_heading('1.2.3 In Scope', level=3)
doc.add_paragraph('User registration and authentication (Passenger/Driver/Admin)', style='List Bullet')
doc.add_paragraph('Ride booking with real-time price estimation', style='List Bullet')
doc.add_paragraph('GPS-based tracking and route monitoring', style='List Bullet')
doc.add_paragraph('Pink Pass enrollment and verification', style='List Bullet')
doc.add_paragraph('Automated safety alerts', style='List Bullet')
doc.add_paragraph('Smart driver matching algorithm', style='List Bullet')
doc.add_paragraph('Admin dashboard with analytics', style='List Bullet')
doc.add_paragraph('Cash payment processing', style='List Bullet')

doc.add_heading('1.2.4 Out of Scope', level=3)
doc.add_paragraph('Multi-city operations (Phase 1 limited to Lahore)', style='List Bullet')
doc.add_paragraph('Third-party payment gateway integration', style='List Bullet')
doc.add_paragraph('Carpooling/ride-sharing features', style='List Bullet')
doc.add_paragraph('Advanced fraud detection systems', style='List Bullet')
doc.add_paragraph('International operations', style='List Bullet')

doc.add_page_break()

doc.add_heading('1.3 Definitions, Acronyms, and Abbreviations', level=2)
definitions_headers = ['Term', 'Definition']
definitions_data = [
    ['AI', 'Artificial Intelligence'],
    ['API', 'Application Programming Interface'],
    ['CNN', 'Convolutional Neural Network'],
    ['CNIC', 'Computerized National Identity Card'],
    ['EAR', 'Eye Aspect Ratio'],
    ['GPS', 'Global Positioning System'],
    ['JWT', 'JSON Web Token [7]'],
    ['MERN', 'MongoDB, Express.js, React, Node.js'],
    ['NLP', 'Natural Language Processing'],
    ['PKR', 'Pakistani Rupees'],
    ['REST', 'Representational State Transfer'],
    ['SMS', 'Short Message Service'],
    ['SOS', 'Emergency distress signal'],
    ['SRS', 'Software Requirements Specification'],
    ['TPR', 'True Positive Rate'],
    ['FPR', 'False Positive Rate']
]
add_table_with_data(doc, definitions_headers, definitions_data, [4, 12])
doc.add_paragraph()

doc.add_heading('1.4 References', level=2)
doc.add_paragraph('This document refers to several international standards, academic studies, and technical documentation. A complete list of all 12 references is provided in Section 7.')

doc.add_heading('1.5 Overview', level=2)
doc.add_paragraph('This SRS document is organized as follows:')
doc.add_paragraph('Section 2 provides an overall description of the SAFORA system, including product perspective, functions, user characteristics, and constraints.', style='List Bullet')
doc.add_paragraph('Section 3 details specific functional and non-functional requirements.', style='List Bullet')
doc.add_paragraph('Section 4 presents system models including use case, sequence, class, and ER diagrams.', style='List Bullet')
doc.add_paragraph('Section 5 outlines the project plan with methodology, timeline, and resource allocation.', style='List Bullet')
doc.add_paragraph('Section 6 contains appendices with supplementary information.', style='List Bullet')

doc.add_page_break()

# ==================== 2. OVERALL DESCRIPTION ====================
doc.add_heading('2. OVERALL DESCRIPTION', level=1)

doc.add_heading('2.1 Product Perspective', level=2)
doc.add_paragraph('SAFORA is a new, self-contained ride-hailing platform designed to address specific safety and efficiency gaps in existing services operating in Pakistan. The system consists of four main components:')

doc.add_heading('2.1.1 System Architecture', level=3)
p = doc.add_paragraph()
p.add_run('Backend Server (Node.js/Express)').font.bold = True
doc.add_paragraph('RESTful API for all client applications', style='List Bullet')
doc.add_paragraph('Real-time communication via Socket.io', style='List Bullet')
doc.add_paragraph('MongoDB database for data persistence', style='List Bullet')
doc.add_paragraph('JWT-based authentication', style='List Bullet')

p = doc.add_paragraph()
p.add_run('AI Microservice (Python/Flask)').font.bold = True
doc.add_paragraph('Pink Pass liveness detection', style='List Bullet')
doc.add_paragraph('Price prediction using Linear Regression', style='List Bullet')
doc.add_paragraph('Smart driver matching algorithm', style='List Bullet')
doc.add_paragraph('Sentiment analysis and demand clustering', style='List Bullet')

p = doc.add_paragraph()
p.add_run('Mobile Applications (React Native) [6]').font.bold = True
doc.add_paragraph('Passenger app for iOS and Android', style='List Bullet')
doc.add_paragraph('Driver app for iOS and Android', style='List Bullet')
doc.add_paragraph('Real-time GPS tracking', style='List Bullet')
doc.add_paragraph('In-app notifications', style='List Bullet')

p = doc.add_paragraph()
p.add_run('Admin Dashboard (React.js)').font.bold = True
doc.add_paragraph('Web-based monitoring interface', style='List Bullet')
doc.add_paragraph('Real-time safety alerts', style='List Bullet')
doc.add_paragraph('Analytics and reporting', style='List Bullet')
doc.add_paragraph('Heatmap visualization', style='List Bullet')

doc.add_heading('2.1.2 System Interfaces', level=3)
p = doc.add_paragraph()
p.add_run('Hardware Interfaces:').font.bold = True
doc.add_paragraph('Mobile device GPS sensors', style='List Bullet')
doc.add_paragraph('Mobile device cameras (for Pink Pass verification)', style='List Bullet')
doc.add_paragraph('Server infrastructure (cloud or on-premise)', style='List Bullet')

p = doc.add_paragraph()
p.add_run('Software Interfaces:').font.bold = True
doc.add_paragraph('Google Maps API for routing and geocoding', style='List Bullet')
doc.add_paragraph('Twilio API for SMS notifications', style='List Bullet')
doc.add_paragraph('MongoDB database (v6.0+)', style='List Bullet')
doc.add_paragraph('Node.js runtime (v18+)', style='List Bullet')
doc.add_paragraph('Python runtime (v3.9+)', style='List Bullet')

p = doc.add_paragraph()
p.add_run('Communication Interfaces:').font.bold = True
doc.add_paragraph('HTTPS for secure API communication', style='List Bullet')
doc.add_paragraph('WebSocket (Socket.io) for real-time updates', style='List Bullet')
doc.add_paragraph('SMS for emergency notifications', style='List Bullet')

doc.add_heading('2.2 Product Functions', level=2)

doc.add_heading('2.2.1 User Management', level=3)
doc.add_paragraph('FR-UM-01: The system shall allow users to register as Passenger, Driver, or Admin with CNIC verification.', style='List Bullet')
doc.add_paragraph('FR-UM-02: The system shall authenticate users using email/phone and password with JWT token generation.', style='List Bullet')
doc.add_paragraph('FR-UM-03: The system shall maintain user profiles with personal information, emergency contacts, and preferences.', style='List Bullet')

doc.add_heading('2.2.2 Pink Pass Feature', level=3)
doc.add_paragraph('FR-PP-01: The system shall allow female passengers to enroll in Pink Pass by submitting a video for liveness verification.', style='List Bullet')
doc.add_paragraph('FR-PP-02: The system shall detect faces using Haar Cascade algorithm with >95% accuracy.', style='List Bullet')
doc.add_paragraph('FR-PP-03: The system shall verify liveness using CNN-based blink detection with >98% accuracy and <5% false positive rate.', style='List Bullet')
doc.add_paragraph('FR-PP-04: The system shall complete Pink Pass verification within 15 seconds.', style='List Bullet')
doc.add_paragraph('FR-PP-05: The system shall only match Pink Pass rides with verified female drivers.', style='List Bullet')

doc.add_heading('2.2.3 Ride Booking', level=3)
doc.add_paragraph('FR-RB-01: The system shall allow passengers to request rides by specifying pickup and dropoff locations.', style='List Bullet')
doc.add_paragraph('FR-RB-02: The system shall calculate estimated price using AI price prediction model (±PKR 30 accuracy).', style='List Bullet')
doc.add_paragraph('FR-RB-03: The system shall match passengers with drivers using smart driver selection algorithm.', style='List Bullet')
doc.add_paragraph('FR-RB-04: The system shall provide real-time ride status updates to both passenger and driver.', style='List Bullet')
doc.add_paragraph('FR-RB-05: The system shall track ride duration, distance, and route in real-time.', style='List Bullet')

doc.add_heading('2.2.4 Safety Sentinel', level=3)
doc.add_paragraph('FR-SS-01: The system shall monitor all active rides using GPS tracking at 5-second intervals.', style='List Bullet')
doc.add_paragraph('FR-SS-02: The system shall calculate point-to-line distance between current location and planned route.', style='List Bullet')
doc.add_paragraph('FR-SS-03: The system shall trigger route deviation alert if distance >500m for >30 seconds [4].', style='List Bullet')
doc.add_paragraph('FR-SS-04: The system shall detect suspicious stops if vehicle stationary for >5 minutes [4].', style='List Bullet')
doc.add_paragraph('FR-SS-05: The system shall send parallel alerts to admin dashboard (Socket.io) and emergency contacts (SMS) within 30 seconds.', style='List Bullet')

doc.add_heading('2.2.5 Driver Management', level=3)
doc.add_paragraph('FR-DM-01: The system shall allow drivers to register with license and vehicle information.', style='List Bullet')
doc.add_paragraph('FR-DM-02: The system shall track driver online/offline status and location.', style='List Bullet')
doc.add_paragraph('FR-DM-03: The system shall calculate driver fairness metric: IdleTime/TotalOnlineTime (last 4 hours).', style='List Bullet')
doc.add_paragraph('FR-DM-04: The system shall maintain driver ratings based on passenger feedback.', style='List Bullet')

doc.add_heading('2.2.6 Admin Dashboard', level=3)
doc.add_paragraph('FR-AD-01: The system shall display real-time safety alerts with severity levels.', style='List Bullet')
doc.add_paragraph('FR-AD-02: The system shall generate demand heatmaps using K-Means clustering.', style='List Bullet')
doc.add_paragraph('FR-AD-03: The system shall analyze passenger feedback using TextBlob sentiment analysis.', style='List Bullet')
doc.add_paragraph('FR-AD-04: The system shall provide analytics on rides, drivers, and safety incidents.', style='List Bullet')

doc.add_heading('2.3 User Characteristics', level=2)
user_headers = ['User Type', 'Technical Expertise', 'Primary Needs', 'Usage Frequency']
user_data = [
    ['Passengers', 'Basic smartphone usage', 'Safe, affordable, reliable transportation', 'Daily to weekly'],
    ['Drivers', 'Basic smartphone and navigation', 'Fair ride distribution, consistent income', 'Daily'],
    ['Administrators', 'Moderate to advanced', 'Monitor safety, manage users, analyze data', 'Daily']
]
add_table_with_data(doc, user_headers, user_data, [3, 4, 5, 4])
doc.add_paragraph()

doc.add_heading('2.4 Constraints', level=2)
doc.add_heading('2.4.1 Regulatory Constraints', level=3)
doc.add_paragraph('Must comply with Pakistan\'s data protection regulations', style='List Bullet')
doc.add_paragraph('CNIC verification required for all users', style='List Bullet')
doc.add_paragraph('SMS notifications must follow PTCL guidelines', style='List Bullet')

doc.add_heading('2.4.2 Technical Constraints', level=3)
doc.add_paragraph('Mobile apps must support Android 8.0+ and iOS 13.0+', style='List Bullet')
doc.add_paragraph('Backend must handle 1000+ concurrent users', style='List Bullet')
doc.add_paragraph('API response time must be <500ms for 95% of requests', style='List Bullet')
doc.add_paragraph('GPS accuracy must be within 10 meters', style='List Bullet')

doc.add_heading('2.4.3 Business Constraints', level=3)
doc.add_paragraph('Phase 1 limited to Lahore city', style='List Bullet')
doc.add_paragraph('Cash-only payments initially', style='List Bullet')
doc.add_paragraph('Development timeline: 22 weeks', style='List Bullet')
doc.add_paragraph('Budget constraints for cloud infrastructure', style='List Bullet')

doc.add_heading('2.5 Assumptions and Dependencies', level=2)
doc.add_heading('2.5.1 Assumptions', level=3)
doc.add_paragraph('Users have smartphones with GPS and camera capabilities', style='List Bullet')
doc.add_paragraph('Reliable internet connectivity (3G/4G/WiFi) available', style='List Bullet')
doc.add_paragraph('Google Maps API will remain accessible and affordable', style='List Bullet')
doc.add_paragraph('Twilio SMS service will maintain 99% uptime', style='List Bullet')

doc.add_heading('2.5.2 Dependencies', level=3)
doc.add_paragraph('MongoDB database availability', style='List Bullet')
doc.add_paragraph('Third-party API services (Google Maps, Twilio)', style='List Bullet')
doc.add_paragraph('Mobile OS platform stability', style='List Bullet')
doc.add_paragraph('Cloud infrastructure reliability', style='List Bullet')

doc.add_page_break()

# ==================== 3. SPECIFIC REQUIREMENTS ====================
doc.add_heading('3. SPECIFIC REQUIREMENTS', level=1)
doc.add_heading('3.1 Functional Requirements', level=2)

doc.add_heading('3.1.1 Authentication and Authorization', level=3)
add_functional_requirement(doc, "FR-3.1.1-01", "User Registration",
                           "System shall allow new users to create accounts",
                           "Name, email, phone, password, CNIC, gender, role",
                           "Validate CNIC format, hash password, generate JWT",
                           "User account created, authentication token", "High")

add_functional_requirement(doc, "FR-3.1.1-02", "User Login",
                           "System shall authenticate existing users",
                           "Email/phone and password",
                           "Verify credentials, generate JWT token",
                           "Authentication token, user profile", "High")

doc.add_heading('3.1.2 Pink Pass Verification', level=3)
add_functional_requirement(doc, "FR-3.1.2-01", "Video Upload",
                           "Female passengers can upload video for verification",
                           "5-10 second video file (max 10MB)",
                           "Extract frames, detect face, verify liveness",
                           "Verification status (approved/rejected), confidence score", "High")

add_functional_requirement(doc, "FR-3.1.2-02", "Liveness Detection",
                           "System shall detect live human face vs. photo/video",
                           "Video frames",
                           "Haar Cascade face detection + CNN blink detection",
                           "Liveness score (0-1), blink count, verification result", "High")

doc.add_heading('3.1.3 Ride Management', level=3)
add_functional_requirement(doc, "FR-3.1.3-01", "Ride Request",
                           "Passengers can request rides",
                           "Pickup location, dropoff location, ride type (standard/pink-pass)",
                           "Calculate route, predict price, match driver",
                           "Estimated price, estimated time, matched driver", "High")

add_functional_requirement(doc, "FR-3.1.3-02", "Driver Matching",
                           "System shall match optimal driver to passenger",
                           "Passenger location, available drivers, ride type",
                           "Calculate weighted scores for all nearby drivers",
                           "Ranked list of drivers, top match selected", "High")

doc.add_heading('3.1.4 Safety Monitoring', level=3)
add_functional_requirement(doc, "FR-3.1.4-01", "Route Deviation Detection",
                           "System shall detect when ride deviates from planned route",
                           "Current GPS location, planned route polyline",
                           "Calculate point-to-line distance using Turf.js",
                           "Deviation distance, alert trigger if threshold exceeded", "Critical")

add_functional_requirement(doc, "FR-3.1.4-02", "Emergency Alert System",
                           "System shall send alerts when safety anomaly detected",
                           "Alert type, ride details, passenger info",
                           "Send Socket.io notification to admin, SMS to emergency contacts",
                           "Alert notifications sent, confirmation receipts", "Critical")

doc.add_heading('3.2 Non-Functional Requirements', level=2)

doc.add_heading('3.2.1 Performance Requirements', level=3)
doc.add_paragraph('NFR-3.2.1-01: Response Time', style='List Bullet')
doc.add_paragraph('API endpoints shall respond within 500ms for 95% of requests', style='List Bullet')
doc.add_paragraph('Pink Pass verification shall complete within 15 seconds [11]', style='List Bullet')
doc.add_paragraph('Driver matching shall complete within 3 seconds [5]', style='List Bullet')
doc.add_paragraph('NFR-3.2.1-02: Throughput', style='List Bullet')
doc.add_paragraph('System shall support 1000+ concurrent users', style='List Bullet')
doc.add_paragraph('System shall handle 100+ simultaneous rides', style='List Bullet')
doc.add_paragraph('Database shall process 10,000+ transactions per hour', style='List Bullet')
doc.add_paragraph('NFR-3.2.1-03: Scalability', style='List Bullet')
doc.add_paragraph('System architecture shall support horizontal scaling', style='List Bullet')
doc.add_paragraph('Database shall support sharding for future growth', style='List Bullet')
doc.add_paragraph('API shall be stateless to enable load balancing', style='List Bullet')

doc.add_heading('3.2.2 Security Requirements', level=3)
doc.add_paragraph('NFR-3.2.2-01: Authentication', style='List Bullet')
doc.add_paragraph('All API endpoints shall require JWT authentication [7]', style='List Bullet')
doc.add_paragraph('Pass­words shall be hashed using bcrypt (cost factor 10)', style='List Bullet')
doc.add_paragraph('JWT tokens shall expire after 24 hours [7]', style='List Bullet')
doc.add_paragraph('NFR-3.2.2-02: Data Protection', style='List Bullet')
doc.add_paragraph('CNIC data shall be encrypted using AES-256', style='List Bullet')
doc.add_paragraph('All API communication shall use HTTPS/TLS 1.3', style='List Bullet')
doc.add_paragraph('Personal data shall comply with Pakistan data protection laws', style='List Bullet')
doc.add_paragraph('NFR-3.2.2-03: Authorization', style='List Bullet')
doc.add_paragraph('Role-based access control (Passenger/Driver/Admin)', style='List Bullet')
doc.add_paragraph('Admin-only access to dashboard and analytics', style='List Bullet')
doc.add_paragraph('Drivers can only access assigned ride details', style='List Bullet')

doc.add_heading('3.2.3 Reliability Requirements', level=3)
doc.add_paragraph('NFR-3.2.3-01: Availability', style='List Bullet')
doc.add_paragraph('System shall maintain 99% uptime during business hours', style='List Bullet')
doc.add_paragraph('Planned maintenance windows shall not exceed 4 hours/month', style='List Bullet')
doc.add_paragraph('Critical safety features shall have 99.9% availability', style='List Bullet')
doc.add_paragraph('NFR-3.2.3-02: Fault Tolerance', style='List Bullet')
doc.add_paragraph('System shall gracefully handle third-party API failures', style='List Bullet')
doc.add_paragraph('SMS alerts shall have fallback notification mechanisms', style='List Bullet')
doc.add_paragraph('Database shall have automated backup every 6 hours', style='List Bullet')

doc.add_page_break()

doc.add_heading('3.2.4 Usability Requirements', level=3)
doc.add_paragraph('NFR-3.2.4-01: User Interface', style='List Bullet')
doc.add_paragraph('Mobile apps shall follow platform design guidelines (Material/iOS)', style='List Bullet')
doc.add_paragraph('Critical actions shall require confirmation', style='List Bullet')
doc.add_paragraph('Error messages shall be clear and actionable', style='List Bullet')
doc.add_paragraph('NFR-3.2.4-02: Accessibility', style='List Bullet')
doc.add_paragraph('Text shall be readable with minimum font size 14px', style='List Bullet')
doc.add_paragraph('Color contrast shall meet WCAG 2.1 AA standards', style='List Bullet')
doc.add_paragraph('UI shall support Urdu and English languages', style='List Bullet')

doc.add_heading('3.2.5 Maintainability Requirements', level=3)
doc.add_paragraph('NFR-3.2.5-01: Code Quality', style='List Bullet')
doc.add_paragraph('Code coverage shall exceed 80% for unit tests', style='List Bullet')
doc.add_paragraph('Code shall follow ESLint and PEP 8 standards', style='List Bullet')
doc.add_paragraph('API documentation shall be maintained using Swagger/OpenAPI', style='List Bullet')
doc.add_paragraph('NFR-3.2.5-02: Monitoring', style='List Bullet')
doc.add_paragraph('System shall log all errors and exceptions', style='List Bullet')
doc.add_paragraph('Performance metrics shall be tracked and visualized', style='List Bullet')
doc.add_paragraph('Safety alerts shall be logged with full audit trail', style='List Bullet')

doc.add_heading('3.3 Interface Requirements', level=2)
doc.add_heading('3.3.1 User Interfaces', level=3)
p = doc.add_paragraph()
p.add_run('Mobile App Screens:').font.bold = True
doc.add_paragraph('Authentication (Login/Register)', style='List Number')
doc.add_paragraph('Home/Dashboard', style='List Number')
doc.add_paragraph('Ride Booking', style='List Number')
doc.add_paragraph('Pink Pass Enrollment', style='List Number')
doc.add_paragraph('Ride Tracking', style='List Number')
doc.add_paragraph('Ride History', style='List Number')
doc.add_paragraph('Profile Management', style='List Number')
doc.add_paragraph('Emergency SOS', style='List Number')

p = doc.add_paragraph()
p.add_run('Admin Dashboard Pages:').font.bold = True
doc.add_paragraph('Real-time Monitoring', style='List Number')
doc.add_paragraph('Safety Alerts', style='List Number')
doc.add_paragraph('User Management', style='List Number')
doc.add_paragraph('Analytics & Reports', style='List Number')
doc.add_paragraph('Heatmap Visualization', style='List Number')

doc.add_heading('3.3.2 Hardware Interfaces', level=3)
doc.add_paragraph('GPS sensor for location tracking', style='List Bullet')
doc.add_paragraph('Camera for Pink Pass verification', style='List Bullet')
doc.add_paragraph('Network interface for API communication', style='List Bullet')

doc.add_heading('3.3.3 Software Interfaces', level=3)
doc.add_paragraph('Google Maps API v3 for routing and geocoding', style='List Bullet')
doc.add_paragraph('Twilio REST API for SMS notifications', style='List Bullet')
doc.add_paragraph('MongoDB database driver', style='List Bullet')
doc.add_paragraph('Socket.io for WebSocket communication', style='List Bullet')

doc.add_heading('3.3.4 Communication Interfaces', level=3)
doc.add_paragraph('HTTPS (TLS 1.3) for API requests', style='List Bullet')
doc.add_paragraph('WebSocket for real-time updates', style='List Bullet')
doc.add_paragraph('SMS for emergency notifications', style='List Bullet')

doc.add_page_break()


# ==================== 4. SYSTEM MODELS ====================
doc.add_heading('4. SYSTEM MODELS', level=1)
doc.add_heading('4.1 Use Case Descriptions', level=2)
doc.add_paragraph('This section provides detailed descriptions for the 30 use cases identified for the SAFORA system, covering Passenger, Driver, and Admin functionalities.')

doc.add_heading('4.1.1 Passenger Use Cases', level=3)
# ... use case tables are below, keeping them ...

add_use_case_table(doc, "UC-01", "Passenger Registration", "Passenger",
                   "Allows a new passenger to create a basic account profile.",
                   "App is installed; Smartphone has internet connectivity.",
                   "1. Passenger selects 'Sign Up'.\n2. Enters Name, Email, and Phone Number.\n3. Creates a secure password.\n4. System validates email format.",
                   "Basic account record created in database.")

add_use_case_table(doc, "UC-02", "Register with CNIC OCR", "Passenger, System",
                   "Verifies identity by extracting data from CNIC using AI OCR.",
                   "UC-01 in progress; Camera permissions granted.",
                   "1. Passenger prompted to take photo of CNIC (Front/Back).\n2. System extracts CNIC number and Gender via OCR.\n3. System validates CNIC with government patterns.\n4. System sends OTP for final phone verification.",
                   "Identity verified; Account fully activated.")

add_use_case_table(doc, "UC-03", "Passenger Login", "Passenger",
                   "Authentication via credentials for secure access.",
                   "Account exists and is verified.",
                   "1. User enters Email/Phone and Password.\n2. System verifies credentials against hashed database records.\n3. System issues 24-hour JWT token.\n4. Redirects to Home screen.",
                   "Access granted to passenger dashboard.")

add_use_case_table(doc, "UC-04", "Update Profile & Emergency Contacts", "Passenger",
                   "Manage personal details and safety contacts for the Safety Sentinel.",
                   "Logged in.",
                   "1. Passenger navigates to 'Profile Settings'.\n2. Updates personal info or password.\n3. Adds/Edits up to 5 Emergency Contacts (Name, Phone).\n4. System saves changes.",
                   "Profile and safety contacts updated.")

add_use_case_table(doc, "UC-05", "Pink Pass Enrollment (Identity Scan)", "Female Passenger",
                   "Initial phase of women-only mode activation involving document check.",
                   "Logged in; Registered as Female via CNIC OCR.",
                   "1. Passenger selects 'Pink Pass' registration.\n2. System confirms gender via stored CNIC metadata.\n3. System requests permission for biometric liveness check.",
                   "System ready for biometric verification.")

add_use_case_table(doc, "UC-06", "Pink Pass Liveness Check", "Female Passenger, AI Service",
                   "AI-driven biometric check using face detection and blink detection.",
                   "UC-05 completed.",
                   "1. App opens camera interface.\n2. Passenger directed to blink/smile at camera.\n3. AI Service calculates Eye Aspect Ratio (EAR).\n4. System confirms 'Liveness' vs 'Photo/Video Spoof'.",
                   "Pink Pass activated for the user.")

add_use_case_table(doc, "UC-07", "Request Ride (Standard/Pink Pass)", "Passenger",
                   "Booking a ride after selecting destination and ride type.",
                   "Logged in; Location services enabled.",
                   "1. Passenger sets destination on map.\n2. System shows route and car categories.\n3. Passenger selects 'Standard' or 'Pink Pass' (if eligible).\n4. Passenger confirms request.",
                   "Ride request broadcasted to matching drivers.")

add_use_case_table(doc, "UC-08", "AI Price Prediction & Estimation", "System, AI Service",
                   "Calculates fair transparent pricing based on demand and traffic.",
                   "Destination set in UC-07.",
                   "1. System sends trip data to AI microservice.\n2. AI Model (Linear Regression) predicts price.\n3. System factors in active traffic and surge multipliers.\n4. Displays transparent 'Estimated Fare' with breakdown.",
                   "Passenger views expected cost before confirming.")

add_use_case_table(doc, "UC-09", "Track Ride & Real-time ETA", "Passenger",
                   "Live monitoring of the driver's approach and trip progress.",
                   "Ride accepted or active.",
                   "1. System displays map with live car movement (5s interval).\n2. Displays ETA in minutes.\n3. Notifies user when driver is at pickup.",
                   "Passenger views live progress on map.")

add_use_case_table(doc, "UC-10", "In-App Communication (Chat)", "Passenger, Driver",
                   "Secure chat between passenger and driver without sharing numbers.",
                   "Ride assigned.",
                   "1. Passenger/Driver opens chat icon.\n2. Sends text message regarding pickup location.\n3. System delivers via Socket.io.\n4. Notifications push to recipient.",
                   "Successful coordination between parties.")

add_use_case_table(doc, "UC-11", "Respond to 'Are you okay?' Safety Prompt", "Passenger",
                   "Responding to automated system checks during anomalies.",
                   "Safety Sentinel detects deviation (UC-26).",
                   "1. System detects deviation and sends push alert with sound.\n2. Passenger prompted: 'Are you okay? Respond within 30s'.\n3. Passenger taps 'I am safe' or 'Help!'.",
                   "Alert either suppressed or escalated.")

add_use_case_table(doc, "UC-12", "Trigger Manual SOS", "Passenger",
                   "Manual emergency trigger in case of immediate threat.",
                   "Ride active.",
                   "1. Passenger taps red SOS button.\n2. System immediately broadcasts location to Admin.\n3. System sends emergency SMS to all registered contacts.",
                   "Emergency protocol activated.")

add_use_case_table(doc, "UC-13", "Ride Feedback & Sentiment Submission", "Passenger",
                   "Providing rating and textual feedback for AI analysis.",
                   "Ride completed.",
                   "1. Passenger rates 1-5 stars.\n2. Passenger writes comments about ride experience.\n3. System processes text through Sentiment Analysis (TextBlob) for Admin flags.",
                   "Sentiment data stored for platform quality control.")

add_use_case_table(doc, "UC-14", "View Personal Ride History", "Passenger",
                   "Accessing previous trips and spending history.",
                   "Logged in.",
                   "1. User navigates to 'Ride History'.\n2. System list past rides with dates and costs.\n3. User can view specific route/driver details for past trips.",
                   "History log displayed.")

doc.add_page_break()

doc.add_heading('4.1.2 Driver Use Cases', level=3)

add_use_case_table(doc, "UC-15", "Driver Registration", "Driver",
                   "Sign-up process for drivers including documentation.",
                   "Valid license and vehicle documents available.",
                   "1. Driver enters personal info.\n2. Uploads License, CNIC, and Vehicle Registration.\n3. Selects car category (Sedan/Mini).\n4. Submits for Admin review.",
                   "Driver account created (Status: Pending Approval).")

add_use_case_table(doc, "UC-16", "Driver PWA Login", "Driver",
                   "Secure access to the driver's Progressive Web App interface.",
                   "Account created and approved.",
                   "1. Driver enters credentials.\n2. System verifies and launches PWA dashboard.\n3. Persistent session established.",
                   "Driver accesses management tools.")

add_use_case_table(doc, "UC-17", "Toggle Availability (Online/Offline)", "Driver",
                   "Managing whether the driver is ready to receive requests.",
                   "Logged in.",
                   "1. Driver toggles 'Go Online'.\n2. System starts background location tracking.\n3. Driver becomes visible in Matching Algorithm pool.",
                   "Status updated in real-time server state.")

add_use_case_table(doc, "UC-18", "Accept/Reject Ride", "Driver, System",
                   "Responding to incoming matched ride requests.",
                   "Driver online; Successfully matched by AI (UC-12).",
                   "1. Driver app rings with pickup details.\n2. Driver accepts request within 15s timer.\n3. System links Driver and Passenger IDs for the ride.",
                   "Ride assigned to the driver.")

add_use_case_table(doc, "UC-19", "Navigate to Pickup/Destination", "Driver",
                   "Using integrated Google Maps for turn-by-turn guidance.",
                   "Ride accepted.",
                   "1. Driver taps 'Start Navigation'.\n2. App redirects/opens Google Maps SDK.\n3. Driver follows route to passenger location.",
                   "Efficient arrival at target location.")

add_use_case_table(doc, "UC-20", "Manage Ride Status", "Driver",
                   "Tracking the physical progression of the passenger journey.",
                   "Ride assigned.",
                   "1. Driver marks 'Arrived' at pickup.\n2. Driver marks 'Start Ride' after passenger boards.\n3. Driver marks 'End Ride' at destination.",
                   "Ride state transitioned in backend.")

add_use_case_table(doc, "UC-21", "View Earnings & Fairness Score", "Driver",
                   "Financial transparency including equitable distribution metrics.",
                   "Logged in.",
                   "1. Driver opens 'Earnings' tab.\n2. System shows daily/weekly revenue.\n3. Displays 'Fairness Score' based on idle vs working time.",
                   "Driver informed of work-life balance and pay.")

add_use_case_table(doc, "UC-22", "Rate Passenger", "Driver",
                   "Providing reciprocal feedback on passenger behavior.",
                   "Ride ended.",
                   "1. Driver submits 1-5 star rating for passenger.\n2. Optional feedback regarding safety/etiquette.\n3. System updates passenger average rating.",
                   "Reciprocal accountability maintained.")

doc.add_page_break()

doc.add_heading('4.1.3 Admin Use Cases', level=3)

add_use_case_table(doc, "UC-23", "Admin Secure Login", "Admin",
                   "Restricted access to the platform command center.",
                   "Admin credentials assigned.",
                   "1. Admin enters credentials.\n2. Multi-factor authentication (optional/planned).\n3. Access granted to monitoring dashboard.",
                   "Secure admin session started.")

add_use_case_table(doc, "UC-24", "Real-time Live Monitoring Map", "Admin",
                   "Overview of all active operations across Lahore.",
                   "Logged in.",
                   "1. Admin opens 'Live Map'.\n2. System renders all active rides as dynamic icons.\n3. Updates positions every 5s via Socket.io.",
                   "Visual situational awareness.")

add_use_case_table(doc, "UC-25", "Safety Sentinel Alert Management", "Admin, System",
                   "Centralized handling of automated safety escalations.",
                   "Anomaly detected (UC-26) or SOS triggered (UC-12).",
                   "1. Dashboard flashes red for new Critical Alert.\n2. Admin clicks alert to see Live GPS and Passenger/Driver info.\n3. Admin initiates contact or protocol.",
                   "Alert incident handled or resolved.")

add_use_case_table(doc, "UC-26", "Route Deviation Monitoring (Auto)", "System, Safety Sentinel",
                   "Automated background logic for path safety.",
                   "Ride is currently in 'Active' state.",
                   "1. System calculates point-to-line distance from planned route.\n2. Threshold of 500m exceeded.\n3. System starts 30s confirmation timer.\n4. Trigger escalation (UC-11/UC-25).",
                   "Autonomous safety monitoring performed.")

add_use_case_table(doc, "UC-27", "Manage User Verification & Suspend", "Admin",
                   "Moderation of accounts to ensure platform integrity.",
                   "LoggedIn.",
                   "1. Admin reviews 'Pending Approval' drivers.\n2. Verifies docs vs photos.\n3. Approves or Suspends users flagged for suspicious sentiment (UC-29).",
                   "User base remains verified and safe.")

add_use_case_table(doc, "UC-28", "Generate AI Demand Heatmaps", "Admin, AI Service",
                   "Business intelligence for optimizing driver distribution.",
                   "Logged in.",
                   "1. Admin selects 'Heatmap' tool.\n2. AI Service performs K-Means clustering on ride request density.\n3. Visualizes peak demand areas on map.",
                   "Data-driven operational insight.")

add_use_case_table(doc, "UC-29", "Analyze Feedback Sentiment Analytics", "Admin, AI Service",
                   "Automated screening of thousands of reviews.",
                   "Logged in.",
                   "1. Dashboard shows 'Negative Sentiment Trends'.\n2. Admin reviews drivers with consistently low sentiment scores.\n3. Automates flagging of high-risk behavior.",
                   "Quality assurance at scale.")

add_use_case_table(doc, "UC-30", "System Reports & Export", "Admin",
                   "Generating statutory and performance reports.",
                   "Logged in.",
                   "1. Admin selects parameters (Date range, Data type).\n2. System generates PDF/CSV of revenue, rides, and safety incidents.\n3. Downloads file for records.",
                   "Audit-ready documentation.")

doc.add_heading('4.2 Sequence Diagrams', level=2)
doc.add_heading('4.2.1 Ride Request Flow', level=3)
doc.add_paragraph('1. Passenger selects pickup/dropoff locations', style='List Number')
doc.add_paragraph('2. System calculates route and price', style='List Number')
doc.add_paragraph('3. System matches optimal driver', style='List Number')
doc.add_paragraph('4. Driver receives notification', style='List Number')
doc.add_paragraph('5. Driver accepts ride', style='List Number')
doc.add_paragraph('6. Passenger receives driver details', style='List Number')
doc.add_paragraph('7. Ride begins with GPS tracking', style='List Number')

doc.add_heading('4.2.2 Safety Alert Flow', level=3)
doc.add_paragraph('1. Safety Sentinel detects route deviation', style='List Number')
doc.add_paragraph('2. System creates alert record', style='List Number')
doc.add_paragraph('3. Admin dashboard receives Socket.io notification', style='List Number')
doc.add_paragraph('4. Emergency contacts receive SMS', style='List Number')
doc.add_paragraph('5. Admin reviews and resolves alert', style='List Number')

doc.add_heading('4.3 Class Diagrams', level=2)
class_headers = ['Class Name', 'Attributes', 'Methods']
class_data = [
    ['User', 'id, name, email, phone, password, role', 'register(), login(), updateProfile()'],
    ['Passenger', 'emergencyContacts, pinkPassVerified', 'enrollPinkPass(), requestRide()'],
    ['Driver', 'licenseNumber, vehicleInfo, rating, location', 'updateStatus(), acceptRide()'],
    ['Ride', 'passenger, driver, pickup, dropoff, status', 'create(), updateStatus(), complete()'],
    ['Alert', 'ride, type, severity, location', 'create(), notify(), resolve()']
]
add_table_with_data(doc, class_headers, class_data, [3, 7, 6])
doc.add_paragraph()

doc.add_heading('4.4 Entity Relationship Diagrams', level=2)
er_headers = ['Entity', 'Attributes', 'Relationships']
er_data = [
    ['Users', '_id, name, email, phone, cnic, gender, role', 'One-to-One with Drivers'],
    ['Drivers', '_id, user, license, vehicle, rating, location', 'One-to-Many with Rides'],
    ['Rides', '_id, passenger, driver, pickup, dropoff, status', 'Many-to-One with Passengers/Drivers'],
    ['Alerts', '_id, ride, type, severity, location', 'Many-to-One with Rides']
]
add_table_with_data(doc, er_headers, er_data, [3, 7, 6])
doc.add_paragraph()

doc.add_page_break()

# ==================== 5. PROJECT PLAN ====================
doc.add_heading('5. PROJECT PLAN', level=1)

doc.add_heading('5.1 Development Methodology', level=2)
p = doc.add_paragraph()
p.add_run('Methodology: Agile Scrum with 2-week sprints [12]').font.bold = True
doc.add_paragraph('Rationale: Agile allows iterative development, frequent testing, and flexibility to adapt to changing requirements.')

p = doc.add_paragraph()
p.add_run('Team Roles:').font.bold = True
doc.add_paragraph('Mirza Umer Ikram: Backend Development, AI Services', style='List Bullet')
doc.add_paragraph('Ruhma Bilal: Frontend Development, UI/UX Design', style='List Bullet')
doc.add_paragraph('Tanveer Ahmed: Project Advisor, Technical Guidance', style='List Bullet')

doc.add_heading('5.2 Timeline and Milestones', level=2)
doc.add_heading('Phase 1: Project Setup (Weeks 1-2) ✓ COMPLETED', level=3)
doc.add_paragraph('Project structure creation', style='List Bullet')
doc.add_paragraph('Development environment setup', style='List Bullet')
doc.add_paragraph('SRS documentation', style='List Bullet')

doc.add_heading('Phase 2: Backend Core (Weeks 3-5) ✓ COMPLETED', level=3)
doc.add_paragraph('Express server with routing', style='List Bullet')
doc.add_paragraph('MongoDB schemas and connection', style='List Bullet')
doc.add_paragraph('JWT authentication system', style='List Bullet')
doc.add_paragraph('User management APIs', style='List Bullet')

doc.add_heading('Phase 3: AI Microservice (Weeks 6-8) ✓ COMPLETED', level=3)
doc.add_paragraph('Flask application structure', style='List Bullet')
doc.add_paragraph('Pink Pass liveness detection', style='List Bullet')
doc.add_paragraph('Price prediction model', style='List Bullet')
doc.add_paragraph('Driver matching algorithm', style='List Bullet')

doc.add_heading('Phase 4: Mobile Apps (Weeks 9-12)', level=3)
doc.add_paragraph('React Native project setup', style='List Bullet')
doc.add_paragraph('Authentication screens', style='List Bullet')
doc.add_paragraph('Ride booking interface', style='List Bullet')
doc.add_paragraph('Real-time tracking UI', style='List Bullet')

doc.add_heading('Phase 5: Safety Features (Weeks 13-15)', level=3)
doc.add_paragraph('GPS tracking system', style='List Bullet')
doc.add_paragraph('Safety Sentinel implementation', style='List Bullet')
doc.add_paragraph('Alert system (Socket.io + Twilio)', style='List Bullet')
doc.add_paragraph('SOS emergency features', style='List Bullet')

doc.add_heading('Phase 6: Admin Dashboard (Weeks 16-17)', level=3)
doc.add_paragraph('React.js dashboard', style='List Bullet')
doc.add_paragraph('Real-time monitoring', style='List Bullet')
doc.add_paragraph('Analytics and reports', style='List Bullet')
doc.add_paragraph('Heatmap visualization', style='List Bullet')

doc.add_heading('Phase 7: AI Model Training & Refinement (Weeks 18-19)', level=3)
doc.add_paragraph('Collect training data', style='List Bullet')
doc.add_paragraph('Train Pink Pass CNN model', style='List Bullet')
doc.add_paragraph('Train price prediction model', style='List Bullet')
doc.add_paragraph('Implement K-Means clustering', style='List Bullet')

doc.add_heading('Phase 8: Testing & Deployment (Weeks 20-22)', level=3)
doc.add_paragraph('Unit testing', style='List Bullet')
doc.add_paragraph('Integration testing', style='List Bullet')
doc.add_paragraph('User acceptance testing', style='List Bullet')
doc.add_paragraph('Production deployment', style='List Bullet')

doc.add_heading('5.3 Resource Allocation', level=2)
doc.add_heading('5.3.1 Human Resources', level=3)
doc.add_paragraph('2 Developers (Full-time)', style='List Bullet')
doc.add_paragraph('1 Project Advisor (Part-time)', style='List Bullet')

doc.add_heading('5.3.2 Technical Resources', level=3)
doc.add_paragraph('Development laptops (2)', style='List Bullet')
doc.add_paragraph('Cloud hosting (AWS/DigitalOcean)', style='List Bullet')
doc.add_paragraph('Third-party APIs (Google Maps, Twilio)', style='List Bullet')
doc.add_paragraph('Testing devices (Android/iOS)', style='List Bullet')

doc.add_heading('5.3.3 Budget Estimate', level=3)
doc.add_paragraph('Cloud infrastructure: PKR 10,000/month', style='List Bullet')
doc.add_paragraph('API costs: PKR 5,000/month', style='List Bullet')
doc.add_paragraph('Testing devices: PKR 50,000 (one-time)', style='List Bullet')
doc.add_paragraph('Miscellaneous: PKR 10,000', style='List Bullet')
p = doc.add_paragraph()
p.add_run('Total Estimated Budget: PKR 140,000').font.bold = True

doc.add_page_break()

# ==================== 6. APPENDICES ====================
doc.add_heading('6. APPENDICES', level=1)
doc.add_heading('6.1 Glossary', level=2)
p = doc.add_paragraph()
p.add_run('Liveness Detection: ').font.bold = True
p.add_run('Biometric technique to determine if a face is from a live person or a photograph/video.')
p = doc.add_paragraph()
p.add_run('Point-to-Line Distance: ').font.bold = True
p.add_run('Geometric calculation measuring shortest distance between a point and a line segment.')
p = doc.add_paragraph()
p.add_run('Weighted Algorithm: ').font.bold = True
p.add_run('Scoring method that assigns different importance levels to multiple factors.')
p = doc.add_paragraph()
p.add_run('Fairness Metric: ').font.bold = True
p.add_run('Quantitative measure ensuring equitable distribution of resources.')

doc.add_heading('6.2 Analysis Models', level=2)
doc.add_heading('6.2.1 AI Model Specifications', level=3)
p = doc.add_paragraph()
p.add_run('Pink Pass CNN Model:').font.bold = True
doc.add_paragraph('Input: Video frames (480x640 RGB)', style='List Bullet')
doc.add_paragraph('Architecture: 3 convolutional layers + 2 dense layers', style='List Bullet')
doc.add_paragraph('Output: Liveness probability (0-1)', style='List Bullet')
doc.add_paragraph('Training data: 10,000+ face images (live vs. spoofed)', style='List Bullet')

p = doc.add_paragraph()
p.add_run('Price Prediction Model:').font.bold = True
doc.add_paragraph('Algorithm: Linear Regression', style='List Bullet')
doc.add_paragraph('Features: 7 (distance, duration, time, day, demand, area, traffic)', style='List Bullet')
doc.add_paragraph('Target: Ride price in PKR', style='List Bullet')
doc.add_paragraph('Expected R²: >0.85', style='List Bullet')

p = doc.add_paragraph()
p.add_run('Driver Matching Algorithm:').font.bold = True
doc.add_paragraph('Scoring formula: 0.50×Distance + 0.30×Rating + 0.20×Fairness', style='List Bullet')
doc.add_paragraph('Distance: Haversine calculation', style='List Bullet')
doc.add_paragraph('Rating: Average passenger ratings (1-5)', style='List Bullet')
doc.add_paragraph('Fairness: IdleTime/TotalOnlineTime (last 4 hours)', style='List Bullet')

doc.add_heading('6.3 Technology Stack', level=2)
tech_headers = ['Component', 'Technology', 'Version']
tech_data = [
    ['Backend', 'Node.js', 'v18+'],
    ['Backend', 'Express.js', 'v4.18'],
    ['Backend', 'MongoDB', 'v6.0 [9]'],
    ['Backend', 'Socket.io', 'v4.7 [10]'],
    ['AI Service', 'Python', 'v3.9+'],
    ['AI Service', 'Flask', 'v3.1'],
    ['AI Service', 'TensorFlow/Keras', 'v2.20 [8]'],
    ['AI Service', 'OpenCV', 'v4.13'],
    ['AI Service', 'scikit-learn', 'v1.8 [8]'],
    ['Mobile Apps', 'React Native', 'v0.72+ [6]'],
    ['Admin Dashboard', 'React.js', 'v18+']
]
add_table_with_data(doc, tech_headers, tech_data, [4, 7, 3])
doc.add_paragraph()

doc.add_page_break()

# ==================== 7. REFERENCES ====================
doc.add_heading('7. REFERENCES', level=1)

references = [
    '[1] IEEE Std 830-1998, "IEEE Recommended Practice for Software Requirements Specifications," Institute of Electrical and Electronics Engineers, 1998.',
    '[2] IEEE Std 29148-2018, "ISO/IEC/IEEE International Standard - Systems and software engineering -- Requirements engineering," IEEE, 2018.',
    '[3] I. Sommerville, Software Engineering, 10th ed., Pearson Education, 2015.',
    '[4] S. Freund, N. Hawelka, and B. Gegov, "Safety and Security Issues in App-Based On-Demand Economy: An Analysis of User Reviews for Ride-Sharing Services," Journal of the Association for Information Systems, vol. 21, no. 4, pp. 1044-1073, 2020.',
    '[5] N. Agatz, A. Erera, M. Savelsbergh, and X. Wang, "Dynamic Ride-Sharing: A Simulation Study in Metro Atlanta," Transportation Research Part B: Methodological, vol. 45, no. 9, pp. 1450-1464, 2011.',
    '[6] D. Goel and P. Jain, "Mobile Application Development: React Native vs Native Android and iOS," International Journal of Computer Science and Mobile Computing, vol. 9, no. 4, pp. 30-35, 2020.',
    '[7] M. Jones, J. Bradley, and N. Sakimura, "JSON Web Token (JWT)," RFC 7519, Internet Engineering Task Force (IETF), May 2015.',
    '[8] A. Géron, Hands-On Machine Learning with Scikit-Learn, Keras, and TensorFlow, 2nd ed., O\'Reilly Media, 2019.',
    '[9] K. Chodorow, MongoDB: The Definitive Guide, 3rd ed., O\'Reilly Media, 2019.',
    '[10] G. Rauch, "Socket.io: Real-time application framework for Node.js," [Online]. Available: https://socket.io. [Accessed: 14-Feb-2026].',
    '[11] A. Anjos and S. Marcel, "Counter-Measures to Photo Attacks in Face Recognition: A Public Database and a Baseline," 2011 International Joint Conference on Biometrics (IJCB), Washington, DC, 2011, pp. 1-7.',
    '[12] K. Schwaber and J. Sutherland, "The Scrum Guide: The Definitive Guide to Scrum: The Rules of the Game," [Online]. Available: https://scrumguides.org. [Accessed: 14-Feb-2026].'
]

for ref in references:
    p = doc.add_paragraph(ref)
    p.paragraph_format.space_after = Pt(6)

# ==================== DOCUMENT APPROVAL ====================
doc.add_paragraph()
doc.add_heading('Document Approval', level=1)
doc.add_paragraph('This Software Requirements Specification has been reviewed and approved by:')
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Project Advisor:\n').font.bold = True
p.add_run('Tanveer Ahmed\n')
p.add_run('Date: _______________')
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Team Members:\n').font.bold = True
p.add_run('Mirza Umer Ikram (S3F22UBSCS081)\n')
p.add_run('Date: _______________')
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Ruhma Bilal (S3F22UBSCS088)\n')
p.add_run('Date: _______________')
doc.add_paragraph()
doc.add_paragraph()

try:
    set_update_fields_on_open(doc)
    print("Enabled automatic field updates on document open.")
except Exception as e:
    print(f"Warning: Could not set updateFields trigger: {e}")

output_path = 'SAFORA_SRS1_IEEE_Format_Complete_Final.docx'
doc.save(output_path)
print(f"Document generated successfully: {output_path}")
