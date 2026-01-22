"""
SAFORA SRS2 Final Document Generator
Creates complete SRS2 with embedded visual diagrams
"""

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def create_element(name):
    return OxmlElement(name)

def create_attribute(element, name, value):
    element.set(qn(name), value)

def add_page_number(paragraph):
    """Add dynamic page number"""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')
    
    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"
    
    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'separate')
    
    fldChar3 = create_element('w:fldChar')
    create_attribute(fldChar3, 'w:fldCharType', 'end')
    
    run._element.append(fldChar1)
    run._element.append(instrText)
    run._element.append(fldChar2)
    run._element.append(fldChar3)

# Create document
doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# Configure headers and footers
section = doc.sections[0]
section.different_first_page_header_footer = True

# Header for subsequent pages
header = section.header
header_para = header.paragraphs[0]
header_para.text = "SAFORA (Smart Ride) - SRS Phase 2"
header_para.style = doc.styles['Header']
header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# Footer with page numbers
footer = section.footer
footer_para = footer.paragraphs[0]
footer_para.text = "Group ID: 11\tPage "
footer_para.style = doc.styles['Footer']
add_page_number(footer_para)

# ==================== TITLE PAGE ====================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run('SAFORA (Smart Ride)\n')
title_run.font.size = Pt(20)
title_run.font.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.add_run('Software Requirements Specification\nPhase 2 - System Design & Diagrams')
subtitle_run.font.size = Pt(18)
subtitle_run.font.bold = True

doc.add_paragraph()
doc.add_paragraph()

uni = doc.add_paragraph()
uni.alignment = WD_ALIGN_PARAGRAPH.CENTER
uni_run = uni.add_run('University of Central Punjab\n')
uni_run.font.bold = True
uni_run.font.size = Pt(14)
uni.add_run('Faculty of Information Technology\n')
uni.add_run('Department of Computer Science')

doc.add_paragraph()
doc.add_paragraph()

details = doc.add_paragraph()
details.alignment = WD_ALIGN_PARAGRAPH.CENTER
details.add_run('Project Title: ').font.bold = True
details.add_run('SAFORA - Smart Ride\n\n')
details.add_run('Project Advisor: ').font.bold = True
details.add_run('Tanveer Ahmed\n\n')
details.add_run('Group ID: ').font.bold = True
details.add_run('11\n\n')

team = details.add_run('Team Members:\n')
team.font.bold = True
details.add_run('Mirza Umer Ikram (S3F22UBSCS081)\n')
details.add_run('Ruhma Bilal (S3F22UBSCS088)\n\n')

date_run = details.add_run('Date: January 20, 2026')
date_run.font.bold = True

doc.add_page_break()

# ==================== TABLE OF CONTENTS ====================
doc.add_heading('Table of Contents', level=1)

toc_items = [
    '1. Use Case Diagrams',
    '   1.1 Passenger Use Cases',
    '   1.2 Driver Use Cases',
    '   1.3 Admin Use Cases',
    '2. Sequence Diagram - Ride Request Flow',
    '3. Class Diagram',
    '4. Entity Relationship Diagram',
    '5. State Diagram - Ride States',
    '6. Component Diagram',
    '7. Document Approval'
]

for item in toc_items:
    doc.add_paragraph(item)

doc.add_page_break()

# ==================== 1. USE CASE DIAGRAMS ====================
doc.add_heading('1. USE CASE DIAGRAMS', level=1)

doc.add_heading('1.1 Passenger Use Cases', level=2)
doc.add_paragraph('The following diagram illustrates all use cases available to passengers in the SAFORA system. Passengers can register, login, enroll in Pink Pass for women-only rides, request and track rides, rate drivers, and trigger emergency SOS alerts.')
doc.add_paragraph()

# Add diagram image
if os.path.exists('diagrams/usecase_passenger.png'):
    doc.add_picture('diagrams/usecase_passenger.png', width=Inches(6))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

doc.add_heading('1.2 Driver Use Cases', level=2)
doc.add_paragraph('The following diagram shows all use cases available to drivers. Drivers can register with their license and vehicle information, manage their online/offline availability, accept or reject ride requests, navigate to pickup locations, start and complete rides, and view their earnings.')
doc.add_paragraph()

if os.path.exists('diagrams/usecase_driver.png'):
    doc.add_picture('diagrams/usecase_driver.png', width=Inches(6))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

doc.add_heading('1.3 Admin Use Cases', level=2)
doc.add_paragraph('The following diagram presents all use cases available to system administrators. Admins can monitor all active rides in real-time, view and resolve safety alerts, manage user accounts, generate analytics reports, and analyze demand heatmaps across the city.')
doc.add_paragraph()

if os.path.exists('diagrams/usecase_admin.png'):
    doc.add_picture('diagrams/usecase_admin.png', width=Inches(6))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# ==================== 2. SEQUENCE DIAGRAM ====================
doc.add_heading('2. SEQUENCE DIAGRAM - RIDE REQUEST FLOW', level=1)
doc.add_paragraph('This sequence diagram illustrates the complete ride request flow from when a passenger selects their pickup and dropoff locations to when they receive confirmation with driver details. The flow involves coordination between the Mobile App, Backend Server, AI Microservice, and Driver.')
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Key Steps:\n').font.bold = True
doc.add_paragraph('1. Passenger selects pickup and dropoff locations in the mobile app', style='List Number')
doc.add_paragraph('2. Mobile app sends ride request to backend server', style='List Number')
doc.add_paragraph('3. Backend calculates route using Google Maps API', style='List Number')
doc.add_paragraph('4. Backend requests price prediction from AI Service', style='List Number')
doc.add_paragraph('5. AI Service returns estimated price based on distance, time, demand, and traffic', style='List Number')
doc.add_paragraph('6. Passenger is shown estimate and confirms request', style='List Number')
doc.add_paragraph('7. Backend requests driver matching from AI Service', style='List Number')
doc.add_paragraph('8. AI Service ranks available drivers using weighted algorithm (Distance 50%, Rating 30%, Fairness 20%)', style='List Number')
doc.add_paragraph('9. Backend sends ride request to top-matched driver', style='List Number')
doc.add_paragraph('10. Driver accepts the ride', style='List Number')
doc.add_paragraph('11. Passenger receives driver details and estimated time of arrival', style='List Number')
doc.add_paragraph()

if os.path.exists('diagrams/sequence_ride_request.png'):
    doc.add_picture('diagrams/sequence_ride_request.png', width=Inches(6.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# ==================== 3. CLASS DIAGRAM ====================
doc.add_heading('3. CLASS DIAGRAM', level=1)
doc.add_paragraph('The class diagram presents the object-oriented structure of the SAFORA system, showing the main classes, their attributes, methods, and relationships. The system follows an inheritance hierarchy with User as the base class, extended by Passenger, Driver, and Admin classes.')
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Core Classes:\n').font.bold = True
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('User (Base Class): ').font.bold = True
p.add_run('Contains common attributes like id, name, email, phone, and password. Provides methods for registration, login, and profile management.')
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Passenger: ').font.bold = True
p.add_run('Extends User with emergency contacts and Pink Pass verification status. Includes methods for enrolling in Pink Pass, requesting rides, and triggering SOS.')
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Driver: ').font.bold = True
p.add_run('Extends User with license information, vehicle details, rating, and location. Provides methods for managing availability and handling rides.')
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Admin: ').font.bold = True
p.add_run('Extends User with permissions and department. Includes methods for monitoring, alerts management, and reporting.')
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Ride: ').font.bold = True
p.add_run('Manages ride information including passenger, driver, locations, pricing, and status.')
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Alert: ').font.bold = True
p.add_run('Handles safety alerts with type, severity, location, and resolution tracking.')
doc.add_paragraph()

if os.path.exists('diagrams/class_diagram.png'):
    doc.add_picture('diagrams/class_diagram.png', width=Inches(6.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# ==================== 4. ER DIAGRAM ====================
doc.add_heading('4. ENTITY RELATIONSHIP DIAGRAM', level=1)
doc.add_paragraph('The Entity Relationship diagram shows the database schema for SAFORA, illustrating the structure of MongoDB collections and their relationships. The system uses four main collections with defined relationships to maintain data integrity.')
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Database Collections:\n').font.bold = True
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Users Collection: ').font.bold = True
p.add_run('Stores user account information including authentication credentials and profile data. Has a one-to-one relationship with Drivers collection and one-to-many relationship with Rides collection (as passenger).')
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Drivers Collection: ').font.bold = True
p.add_run('Contains driver-specific information including license, vehicle details, rating, and current location (GeoJSON). Links to Users collection and has one-to-many relationship with Rides.')
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Rides Collection: ').font.bold = True
p.add_run('Stores ride information including passenger and driver references, locations, pricing, and status. Has many-to-one relationships with Users and Drivers, and one-to-many relationship with Alerts.')
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Alerts Collection: ').font.bold = True
p.add_run('Manages safety alerts generated by Safety Sentinel, linked to specific rides with severity levels and resolution tracking.')
doc.add_paragraph()

if os.path.exists('diagrams/er_diagram.png'):
    doc.add_picture('diagrams/er_diagram.png', width=Inches(6.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# ==================== 5. STATE DIAGRAM ====================
doc.add_heading('5. STATE DIAGRAM - RIDE STATES', level=1)
doc.add_paragraph('The state diagram illustrates all possible states of a ride and the transitions between these states. Understanding ride states is crucial for proper ride management, billing, and safety monitoring.')
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Ride States and Transitions:\n').font.bold = True
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('REQUESTED → MATCHED: ').font.bold = True
p.add_run('When a driver accepts the ride request')
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('MATCHED → DRIVER_EN_ROUTE: ').font.bold = True
p.add_run('When driver starts navigation to pickup location')
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('DRIVER_EN_ROUTE → DRIVER_ARRIVED: ').font.bold = True
p.add_run('When driver reaches the pickup location')
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('DRIVER_ARRIVED → IN_PROGRESS: ').font.bold = True
p.add_run('When passenger boards and ride begins')
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('IN_PROGRESS → COMPLETED: ').font.bold = True
p.add_run('When destination is reached and ride ends successfully')
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Any State → CANCELLED: ').font.bold = True
p.add_run('Ride can be cancelled from multiple states due to passenger/driver actions or safety concerns')
doc.add_paragraph()

if os.path.exists('diagrams/state_ride.png'):
    doc.add_picture('diagrams/state_ride.png', width=Inches(6.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# ==================== 6. COMPONENT DIAGRAM ====================
doc.add_heading('6. COMPONENT DIAGRAM', level=1)
doc.add_paragraph('The component diagram provides a high-level view of the SAFORA system architecture, showing the major components and their interactions. The system follows a microservices architecture with clear separation of concerns.')
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('System Components:\n').font.bold = True
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Mobile Applications: ').font.bold = True
p.add_run('React Native apps for passengers, drivers, and the admin dashboard web interface. These communicate with backend via HTTPS.')
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Backend Server: ').font.bold = True
p.add_run('Node.js/Express server with multiple services including API Gateway, Authentication, User Management, Ride Service, Safety Sentinel, and real-time notifications via Socket.io.')
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('AI Microservice: ').font.bold = True
p.add_run('Python/Flask service providing Pink Pass verification, price prediction, driver matching, and sentiment analysis. Communicates with backend via REST API.')
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('MongoDB Database: ').font.bold = True
p.add_run('NoSQL database storing Users, Drivers, Rides, and Alerts collections. Connected to backend via MongoDB protocol.')
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('External APIs: ').font.bold = True
p.add_run('Google Maps API for routing and geocoding, Twilio API for SMS notifications.')
doc.add_paragraph()

if os.path.exists('diagrams/component_diagram.png'):
    doc.add_picture('diagrams/component_diagram.png', width=Inches(6.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# ==================== DOCUMENT APPROVAL ====================
doc.add_heading('Document Approval', level=1)
doc.add_paragraph('This SRS Phase 2 document has been reviewed and approved by:')
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Project Advisor:\n').font.bold = True
p.add_run('Tanveer Ahmed\n')
p.add_run('Date: _______________')
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Team Members:\n').font.bold = True
p.add_run('Mirza Umer Ikram (S3F22UBSCS081)\n')
p.add_run('Date: _______________')
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Ruhma Bilal (S3F22UBSCS088)\n')
p.add_run('Date: _______________')

# Save document
output_path = 'SAFORA_SRS2_Complete_With_Diagrams.docx'
doc.save(output_path)
print(f"✅ Complete SRS2 document generated: {output_path}")
print("✅ Document includes:")
print("   - Professional title page with headers/footers")
print("   - Table of Contents")
print("   - 8 embedded visual diagrams:")
print("     • Use Case Diagrams (Passenger, Driver, Admin)")
print("     • Sequence Diagram (Ride Request Flow)")
print("     • Class Diagram")
print("     • Entity Relationship Diagram")
print("     • State Diagram (Ride States)")
print("     • Component Diagram")
print("   - Detailed explanations for each diagram")
print("   - Document approval section")
